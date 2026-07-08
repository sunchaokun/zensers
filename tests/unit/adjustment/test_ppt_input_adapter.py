import pytest

from src.core.adjustment.ppt_input_adapter import DocxDataParser, PptInputAdapter, PdfDataParser, ExcelDataParser, TextDataParser, CsvDataParser, JsonDataParser
from src.core.adjustment.extraction_types import ExtractionResult
from src.content.content_orchestrator import ContentSection, SectionType


def _create_docx_with_headings(path, paragraphs):
    from docx import Document
    doc = Document()
    for style, text in paragraphs:
        doc.add_paragraph(text, style=style)
    doc.save(path)


def _create_docx_with_tables(path, rows_data):
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    for i, row in enumerate(rows_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
    doc.save(path)


class TestDocxDataParser:
    def test_parse_with_headings_creates_sections(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_headings(docx_path, [
            ("Heading 1", "Introduction"),
            ("Normal", "This is the intro paragraph."),
            ("Heading 1", "Market Analysis"),
            ("Normal", "The market is growing."),
        ])
        parser = DocxDataParser()
        result = parser.parse(docx_path)
        assert isinstance(result, ExtractionResult)
        assert len(result.sections) == 2
        assert result.sections[0].title == "Introduction"
        assert result.sections[1].title == "Market Analysis"
        assert "intro paragraph" in result.sections[0].points[0]

    def test_parse_without_headings_creates_single_section(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_headings(docx_path, [
            ("Normal", "First paragraph."),
            ("Normal", "Second paragraph."),
        ])
        parser = DocxDataParser()
        result = parser.parse(docx_path)
        assert len(result.sections) == 1
        assert result.sections[0].title == ""
        assert "First paragraph" in result.sections[0].content

    def test_parse_extracts_tables(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_tables(docx_path, [
            ["Name", "Value"],
            ["Revenue", "$10B"],
        ])
        parser = DocxDataParser()
        result = parser.parse(docx_path)
        assert len(result.tables) == 1
        assert result.tables[0][0] == ["Name", "Value"]
        assert result.tables[0][1] == ["Revenue", "$10B"]

    def test_parse_metadata_includes_format(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_headings(docx_path, [("Normal", "text")])
        parser = DocxDataParser()
        result = parser.parse(docx_path)
        assert result.metadata["format"] == "docx"

    def test_parse_nonexistent_file_raises(self):
        parser = DocxDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.docx")

    def test_parse_empty_docx(self, tmp_path):
        docx_path = str(tmp_path / "empty.docx")
        from docx import Document
        Document().save(docx_path)
        parser = DocxDataParser()
        result = parser.parse(docx_path)
        assert result.sections == []
        assert result.tables == []
        assert result.title == ""


class TestPptInputAdapter:
    def test_extract_single_docx(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_headings(docx_path, [
            ("Heading 1", "Title"),
            ("Normal", "Content here."),
        ])
        adapter = PptInputAdapter()
        result = adapter.extract([docx_path])
        assert result.title == "Title"
        assert len(result.sections) == 1

    def test_extract_multiple_docx_merges(self, tmp_path):
        path1 = str(tmp_path / "doc1.docx")
        path2 = str(tmp_path / "doc2.docx")
        _create_docx_with_headings(path1, [("Heading 1", "Part One"), ("Normal", "A")])
        _create_docx_with_headings(path2, [("Heading 1", "Part Two"), ("Normal", "B")])
        adapter = PptInputAdapter()
        result = adapter.extract([path1, path2])
        assert len(result.sections) == 2
        assert result.metadata["file_count"] == 2

    def test_extract_empty_list_returns_empty(self):
        adapter = PptInputAdapter()
        result = adapter.extract([])
        assert result.title == ""
        assert result.sections == []

    def test_extract_unsupported_format_fallback(self, tmp_path):
        unsupported = str(tmp_path / "data.xyz")
        with open(unsupported, "w") as f:
            f.write("some data")
        adapter = PptInputAdapter()
        result = adapter.extract([unsupported])
        assert result.metadata["format"] == "unknown"
        assert result.title == "data"


def _create_pdf_with_text(path, text_pages):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(path, pagesize=letter)
    for page_text in text_pages:
        c.drawString(72, 720, page_text)
        c.showPage()
    c.save()


class TestPdfDataParser:
    def test_parse_extracts_text(self, tmp_path):
        pdf_path = str(tmp_path / "test.pdf")
        _create_pdf_with_text(pdf_path, ["Introduction page", "Market data page"])
        parser = PdfDataParser()
        result = parser.parse(pdf_path)
        assert isinstance(result, ExtractionResult)
        assert result.metadata["format"] == "pdf"
        assert result.metadata["page_count"] == 2

    def test_parse_nonexistent_file_raises(self):
        parser = PdfDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.pdf")

    def test_parse_empty_pdf(self, tmp_path):
        pdf_path = str(tmp_path / "empty.pdf")
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.showPage()
        c.save()
        parser = PdfDataParser()
        result = parser.parse(pdf_path)
        assert result.metadata["format"] == "pdf"


def _create_xlsx_with_data(path, sheets_data):
    import openpyxl
    wb = openpyxl.Workbook()
    first = True
    for name, rows in sheets_data.items():
        if first:
            ws = wb.active
            ws.title = name
            first = False
        else:
            ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    wb.save(path)


class TestExcelDataParser:
    def test_parse_creates_section_per_sheet(self, tmp_path):
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_xlsx_with_data(xlsx_path, {
            "Revenue": [["Year", "Amount"], ["2023", "$10B"], ["2024", "$12B"]],
            "Growth": [["Metric", "Value"], ["CAGR", "5%"]],
        })
        parser = ExcelDataParser()
        result = parser.parse(xlsx_path)
        assert isinstance(result, ExtractionResult)
        assert len(result.sections) == 2
        assert result.sections[0].title == "Revenue"
        assert result.sections[1].title == "Growth"

    def test_parse_extracts_tables(self, tmp_path):
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_xlsx_with_data(xlsx_path, {
            "Data": [["A", "B"], ["1", "2"]],
        })
        parser = ExcelDataParser()
        result = parser.parse(xlsx_path)
        assert len(result.tables) == 1
        assert result.tables[0][0] == ["A", "B"]

    def test_parse_nonexistent_file_raises(self):
        parser = ExcelDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.xlsx")

    def test_parse_empty_workbook(self, tmp_path):
        xlsx_path = str(tmp_path / "empty.xlsx")
        import openpyxl
        wb = openpyxl.Workbook()
        wb.save(xlsx_path)
        parser = ExcelDataParser()
        result = parser.parse(xlsx_path)
        assert result.metadata["format"] == "xlsx"


class TestTextDataParser:
    def test_parse_creates_single_section(self, tmp_path):
        txt_path = str(tmp_path / "test.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("This is a test document.\nIt has two lines.")
        parser = TextDataParser()
        result = parser.parse(txt_path)
        assert len(result.sections) == 1
        assert "test document" in result.sections[0].content
        assert result.metadata["format"] == "txt"

    def test_parse_nonexistent_file_raises(self):
        parser = TextDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.txt")

    def test_parse_empty_file(self, tmp_path):
        txt_path = str(tmp_path / "empty.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("")
        parser = TextDataParser()
        result = parser.parse(txt_path)
        assert result.sections == []


class TestCsvDataParser:
    def test_parse_extracts_headers_and_rows(self, tmp_path):
        csv_path = str(tmp_path / "test.csv")
        with open(csv_path, "w", encoding="utf-8", newline="") as f:
            import csv
            writer = csv.writer(f)
            writer.writerow(["Name", "Value"])
            writer.writerow(["Revenue", "$10B"])
            writer.writerow(["Growth", "5%"])
        parser = CsvDataParser()
        result = parser.parse(csv_path)
        assert len(result.tables) == 1
        assert result.tables[0][0] == ["Name", "Value"]
        assert result.metadata["format"] == "csv"

    def test_parse_nonexistent_file_raises(self):
        parser = CsvDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.csv")


class TestJsonDataParser:
    def test_parse_array_of_objects(self, tmp_path):
        json_path = str(tmp_path / "test.json")
        import json
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump([
                {"name": "Revenue", "value": "$10B"},
                {"name": "Growth", "value": "5%"},
            ], f)
        parser = JsonDataParser()
        result = parser.parse(json_path)
        assert len(result.sections) == 2
        assert result.metadata["format"] == "json"

    def test_parse_nested_object(self, tmp_path):
        json_path = str(tmp_path / "nested.json")
        import json
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "title": "Report",
                "sections": [
                    {"heading": "Intro", "text": "Hello"},
                    {"heading": "Data", "text": "Numbers"},
                ],
            }, f)
        parser = JsonDataParser()
        result = parser.parse(json_path)
        assert result.title == "Report"
        assert result.metadata["format"] == "json"

    def test_parse_nonexistent_file_raises(self):
        parser = JsonDataParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.json")
