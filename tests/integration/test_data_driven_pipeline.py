import os
import json
import pytest
from pathlib import Path

from src.core.adjustment.ppt_input_adapter import (
    PptInputAdapter, DocxDataParser, PdfDataParser,
    ExcelDataParser, TextDataParser, CsvDataParser, JsonDataParser,
)
from src.core.adjustment.extraction_types import ExtractionResult, ExtractionSummary
from src.core.adjustment.ppt_requirement_extractor import PptRequirementExtractor, PptRequirement
from src.core.adjustment.ppt_data_supplementer import PptDataSupplementer, DataGap
from src.core.adjustment.slide_data_builder import SlideDataBuilder
from src.core.dialogue.state_machine import ConversationStateMachine, ConversationState
from src.content.content_orchestrator import ContentSection, SectionType


class TestFullPipelineDocx:
    def test_docx_to_ppt_full_flow(self, tmp_path):
        from docx import Document
        docx_path = str(tmp_path / "market_report.docx")
        doc = Document()
        doc.add_paragraph("2024年新能源市场分析报告", style="Heading 1")
        doc.add_paragraph("全球新能源市场规模达到$500B，同比增长15%。", style="Normal")
        doc.add_paragraph("竞争格局", style="Heading 1")
        doc.add_paragraph("TOP3企业市场份额超过60%。", style="Normal")
        doc.add_paragraph("特斯拉占据25%份额。", style="Normal")
        doc.add_paragraph("技术趋势", style="Heading 1")
        doc.add_paragraph("固态电池技术突破将改变行业格局。", style="Normal")
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "企业"
        table.rows[0].cells[1].text = "份额"
        table.rows[1].cells[0].text = "特斯拉"
        table.rows[1].cells[1].text = "25%"
        table.rows[2].cells[0].text = "比亚迪"
        table.rows[2].cells[1].text = "18%"
        doc.save(docx_path)

        adapter = PptInputAdapter()
        extraction = adapter.extract([docx_path])

        assert extraction.title == "2024年新能源市场分析报告"
        assert len(extraction.sections) == 3
        assert len(extraction.tables) == 1
        assert extraction.tables[0][1] == ["特斯拉", "25%"]
        assert "竞争格局" in extraction.key_topics
        assert "技术趋势" in extraction.key_topics

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction, user_description="做一个关于新能源市场的汇报PPT")
        assert "新能源" in requirement.topic
        assert requirement.page_count >= 3
        assert len(requirement.focus) > 0

        supplementer = PptDataSupplementer()
        gaps = supplementer.analyze_gaps(extraction, requirement)
        for gap in gaps:
            assert gap.topic not in extraction.key_topics
            assert len(gap.search_queries) > 0

        builder = SlideDataBuilder()
        slide_data_list = builder.build_list(extraction.sections, add_cover=True, add_end=True, title=requirement.topic)
        assert len(slide_data_list) >= 5
        assert slide_data_list[0]["slide_type"] == "cover"
        assert slide_data_list[-1]["slide_type"] == "end"
        for sd in slide_data_list[1:-1]:
            assert "title" in sd
            assert "items" in sd

    def test_docx_no_headings_to_ppt(self, tmp_path):
        from docx import Document
        docx_path = str(tmp_path / "plain.docx")
        doc = Document()
        doc.add_paragraph("这是一份纯文本文档。")
        doc.add_paragraph("没有标题结构。")
        doc.save(docx_path)

        adapter = PptInputAdapter()
        extraction = adapter.extract([docx_path])
        assert len(extraction.sections) == 1
        assert "纯文本文档" in extraction.sections[0].content

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction)
        assert requirement.page_count >= 3


class TestFullPipelinePdf:
    def test_pdf_to_ppt_full_flow(self, tmp_path):
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        pdf_path = str(tmp_path / "report.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(72, 720, "Annual Report 2024")
        c.drawString(72, 700, "Revenue: $10B, Growth: 15%")
        c.showPage()
        c.drawString(72, 720, "Market Analysis")
        c.drawString(72, 700, "Competition is intensifying across all segments.")
        c.showPage()
        c.save()

        adapter = PptInputAdapter()
        extraction = adapter.extract([pdf_path])

        assert extraction.metadata["format"] == "pdf"
        assert extraction.metadata["page_count"] == 2
        assert len(extraction.sections) >= 1

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction)
        assert requirement.page_count >= 3


class TestFullPipelineExcel:
    def test_excel_to_ppt_full_flow(self, tmp_path):
        import openpyxl
        xlsx_path = str(tmp_path / "data.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Revenue"
        ws.append(["Year", "Revenue", "Growth"])
        ws.append(["2023", "$10B", "12%"])
        ws.append(["2024", "$12B", "20%"])
        ws2 = wb.create_sheet("Market Share")
        ws2.append(["Company", "Share"])
        ws2.append(["Tesla", "25%"])
        ws2.append(["BYD", "18%"])
        wb.save(xlsx_path)

        adapter = PptInputAdapter()
        extraction = adapter.extract([xlsx_path])

        assert len(extraction.sections) == 2
        assert extraction.sections[0].title == "Revenue"
        assert extraction.sections[1].title == "Market Share"
        assert len(extraction.tables) == 2
        assert extraction.tables[0][0] == ["Year", "Revenue", "Growth"]

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction)
        assert "Revenue" in requirement.focus
        assert "Market Share" in requirement.focus


class TestFullPipelineMultiFile:
    def test_multi_file_merge(self, tmp_path):
        from docx import Document
        docx_path = str(tmp_path / "report.docx")
        doc = Document()
        doc.add_paragraph("市场概述", style="Heading 1")
        doc.add_paragraph("市场规模持续增长。", style="Normal")
        doc.save(docx_path)

        txt_path = str(tmp_path / "notes.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("补充数据：\n2024年增长率达到20%\n主要驱动力来自亚太地区")

        adapter = PptInputAdapter()
        extraction = adapter.extract([docx_path, txt_path])

        assert extraction.metadata["file_count"] == 2
        assert len(extraction.sections) == 2
        assert "市场概述" in [s.title for s in extraction.sections]

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction)
        assert requirement.page_count >= 3


class TestFullPipelineStateTransition:
    def test_state_flow_data_extracted_to_framework(self):
        sm = ConversationStateMachine()
        sm.transition(ConversationState.DATA_EXTRACTED)
        assert sm.current_state == ConversationState.DATA_EXTRACTED

        sm.transition(ConversationState.REQUIREMENT_CONFIRM)
        assert sm.current_state == ConversationState.REQUIREMENT_CONFIRM

        sm.transition(ConversationState.DATA_SUPPLEMENT)
        assert sm.current_state == ConversationState.DATA_SUPPLEMENT

        sm.transition(ConversationState.FRAMEWORK_CONFIRM)
        assert sm.current_state == ConversationState.FRAMEWORK_CONFIRM

    def test_data_extracted_can_route_to_executing(self):
        sm = ConversationStateMachine()
        sm.transition(ConversationState.DATA_EXTRACTED)
        sm.transition(ConversationState.EXECUTING)
        assert sm.current_state == ConversationState.EXECUTING


class TestFullPipelineSupplementer:
    def test_supplementer_fills_gap_with_mock(self):
        extraction = ExtractionResult(
            title="新能源报告", sections=[], tables=[],
            key_topics=["market size", "competition"],
            metadata={}, summary=None,
        )
        requirement = PptRequirement(
            topic="新能源市场",
            focus=["market size", "competition", "technology", "policy"],
        )
        supplementer = PptDataSupplementer()
        gaps = supplementer.analyze_gaps(extraction, requirement)
        assert len(gaps) == 2
        gap_topics = [g.topic for g in gaps]
        assert "technology" in gap_topics
        assert "policy" in gap_topics

        class MockSearch:
            def execute(self, **kwargs):
                return {"success": True, "data": {"results": [f"Data for {kwargs.get('query', '')}"]}}

        filled = supplementer.supplement(gaps, search_skill=MockSearch())
        assert all(g.filled for g in filled)


class TestFullPipelineCsvJson:
    def test_csv_full_flow(self, tmp_path):
        csv_path = str(tmp_path / "sales.csv")
        import csv
        with open(csv_path, "w", encoding="utf-8", newline="") as f:
            w = csv.writer(f)
            w.writerow(["Region", "Sales", "YoY"])
            w.writerow(["APAC", "$5B", "+18%"])
            w.writerow(["EMEA", "$3B", "+8%"])

        adapter = PptInputAdapter()
        extraction = adapter.extract([csv_path])
        assert len(extraction.tables) == 1
        assert extraction.tables[0][0] == ["Region", "Sales", "YoY"]
        assert "Region" in extraction.key_topics

    def test_json_full_flow(self, tmp_path):
        json_path = str(tmp_path / "config.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "title": "Product Analysis",
                "sections": [
                    {"heading": "Product A", "text": "Revenue $2B"},
                    {"heading": "Product B", "text": "Revenue $1.5B"},
                ],
            }, f)

        adapter = PptInputAdapter()
        extraction = adapter.extract([json_path])
        assert extraction.title == "Product Analysis"
        assert len(extraction.sections) == 2
        assert extraction.sections[0].title == "Product A"


class TestFullPipelineSlideDataBuild:
    def test_extraction_to_slide_data_to_outline(self):
        sections = [
            ContentSection(id="s0", title="市场规模", content="全球市场$500B",
                           order=0, type=SectionType.BODY, points=["CAGR 15%", "亚太领先"]),
            ContentSection(id="s1", title="竞争格局", content="TOP3占60%",
                           order=1, type=SectionType.BODY, points=["特斯拉25%", "比亚迪18%"]),
            ContentSection(id="s2", title="技术趋势", content="固态电池突破",
                           order=2, type=SectionType.BODY, points=["2025量产", "成本降30%"]),
        ]
        extraction = ExtractionResult(
            title="新能源报告", sections=sections, tables=[],
            key_topics=["市场规模", "竞争格局", "技术趋势"],
            metadata={"format": "docx", "page_count": 5}, summary=None,
        )

        extractor = PptRequirementExtractor()
        requirement = extractor.extract(extraction)
        assert requirement.topic == "新能源报告"
        assert requirement.page_count == 6

        builder = SlideDataBuilder()
        slide_data_list = builder.build_list(sections, add_cover=True, add_end=True, title=requirement.topic)
        assert len(slide_data_list) == 5
        assert slide_data_list[0]["slide_type"] == "cover"
        assert slide_data_list[0]["title"] == "新能源报告"
        assert slide_data_list[1]["title"] == "市场规模"
        assert slide_data_list[1]["items"] == ["CAGR 15%", "亚太领先"]
        assert slide_data_list[2]["title"] == "竞争格局"
        assert slide_data_list[-1]["slide_type"] == "end"

        from src.converters.slide_outline_builder import SlideOutlineBuilder
        outline_builder = SlideOutlineBuilder()
        outline = outline_builder.build(slide_data_list)
        assert len(outline.slides) == 5
