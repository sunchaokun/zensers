# -*- coding: utf-8 -*-
"""
HTML to Word Converter
======================

Converts HTML intermediate format to Word documents (.docx).

Core features:
1. HTML parsing and conversion
2. Heading level handling
3. Paragraph and list processing
4. Table processing
5. Style application

Technical solution: python-docx

Usage example:
    converter = HTMLToWordConverter()
    result = converter.convert(
        html="<article><h1>Title</h1><p>Content</p></article>",
        output_path="output.docx"
    )
"""

import logging
import os
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from .base_parser import HTMLElementParser
from .css_extractor import CSSStyleExtractor, ExtractedStyles

logger = logging.getLogger(__name__)

# Constants
MAX_HTML_SIZE = 50 * 1024 * 1024  # 50MB - Maximum HTML content size


@dataclass
class ConversionResult:
    """Conversion result"""
    success: bool
    output_path: Optional[str] = None
    file_size: Optional[int] = None
    pages_estimate: Optional[int] = None
    error: Optional[str] = None
    error_code: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "success": self.success,
            "output_path": self.output_path,
            "file_size": self.file_size,
            "pages_estimate": self.pages_estimate,
            "error": self.error,
            "error_code": self.error_code
        }


class HTMLToWordConverter:
    """
    HTML to Word Converter
    
    Converts HTML intermediate format to Word documents.
    
    Usage example:
        converter = HTMLToWordConverter()
        result = converter.convert(
            html="<article><h1>Title</h1><p>Content</p></article>",
            output_path="output.docx"
        )
        
        if result.success:
            print(f"File saved: {result.output_path}")
    """
    
    # Default styles
    DEFAULT_STYLES = {
        "title_font": "Microsoft YaHei",
        "body_font": "Microsoft YaHei",
        "title_size": 28,
        "h1_size": 24,
        "h2_size": 20,
        "h3_size": 16,
        "body_size": 12,
        "line_spacing": 1.5
    }
    
    def __init__(self, styles: Optional[Dict[str, Any]] = None):
        """
        Initialize converter
        
        Args:
            styles: Custom style configuration
        """
        self.styles = {**self.DEFAULT_STYLES, **(styles or {})}
        self._docx_available = self._check_docx_available()
        
        if not self._docx_available:
            logger.warning("python-docx not available, converter will have limited functionality")
    
    def _check_docx_available(self) -> bool:
        """Check if python-docx is available"""
        try:
            from docx import Document  # noqa: F401
            return True
        except ImportError:
            return False
    
    def get_default_styles(self) -> Dict[str, Any]:
        """Get default styles"""
        return self.DEFAULT_STYLES.copy()
    
    def _merge_styles(
        self,
        extracted_styles: Optional[ExtractedStyles],
        custom_styles: Optional[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        Merge styles
        
        Merge order: Template styles → Default styles → Custom styles
        
        Args:
            extracted_styles: Styles extracted from template
            custom_styles: User custom styles
            
        Returns:
            Merged style dictionary
        """
        # Default styles as base
        final_styles = self.DEFAULT_STYLES.copy()
        
        # Apply template extracted styles (higher priority than defaults)
        if extracted_styles:
            template_styles = extracted_styles.to_word_styles()
            final_styles.update(template_styles)
            logger.debug(f"Applied template styles: {template_styles}")
        
        # Apply custom styles (highest priority)
        if custom_styles:
            final_styles.update(custom_styles)
            logger.debug(f"Applied custom styles: {custom_styles}")
        
        return final_styles
    
    def convert(
        self,
        html: str,
        output_path: str,
        styles: Optional[Dict[str, Any]] = None,
        template_html: Optional[str] = None
    ) -> ConversionResult:
        """
        Convert HTML to Word document
        
        Args:
            html: HTML content
            output_path: Output file path
            styles: Custom styles (optional)
            template_html: HTML template content (for extracting CSS styles)
            
        Returns:
            ConversionResult conversion result
        """
        # ========== CONVERSION START ==========
        logger.info("[CONVERT] ========== Starting HTML to Word conversion ==========")
        
        # Input validation
        if not isinstance(html, str):
            logger.warning("[CONVERT] html is not a string, converting to empty string")
            html = ""
        
        logger.info(f"[CONVERT] Input HTML size: {len(html)} bytes")
        
        if not isinstance(output_path, str):
            logger.error("[CONVERT] FAILED: output_path is not a string")
            return ConversionResult(
                success=False,
                error="output_path must be a string",
                error_code="INVALID_PATH_TYPE"
            )
        
        logger.info(f"[CONVERT] Output path: {output_path}")
        
        # File size limit check
        if len(html) > MAX_HTML_SIZE:
            logger.error(f"[CONVERT] FAILED: HTML too large ({len(html)/1024/1024:.1f}MB > {MAX_HTML_SIZE/1024/1024}MB)")
            return ConversionResult(
                success=False,
                error=f"HTML content too large ({len(html)/1024/1024:.1f}MB > {MAX_HTML_SIZE/1024/1024}MB)",
                error_code="CONTENT_TOO_LARGE"
            )
        
        # Path safety check
        if not self._is_safe_path(output_path):
            logger.error(f"[CONVERT] FAILED: Unsafe path detected: {output_path}")
            return ConversionResult(
                success=False,
                error="Invalid or unsafe output path",
                error_code="UNSAFE_PATH"
            )
        
        logger.info("[CONVERT] Path safety check passed")
        
        # Empty HTML handling
        if not html.strip():
            logger.warning("[CONVERT] Empty HTML, creating minimal document")
            html = "<article><p></p></article>"
        
        # Preprocessing: Remove style blocks and residual template tags
        logger.info("[CONVERT] Step 1: Sanitizing HTML")
        html = self._sanitize_html(html)
        logger.info(f"[CONVERT] Sanitized HTML size: {len(html)} bytes")
        
        # Extract CSS styles from template HTML (if provided)
        extracted_styles = None
        if template_html:
            logger.info("[CONVERT] Step 2: Extracting CSS styles from template")
            extractor = CSSStyleExtractor()
            extracted_styles = extractor.extract_from_html(template_html)
            logger.info(f"[CONVERT] CSS extraction complete: title_size={extracted_styles.title_size}, h1_size={extracted_styles.h1_size}, h2_size={extracted_styles.h2_size}")
        else:
            logger.info("[CONVERT] Step 2: No template HTML, using default styles")
        
        # Merge styles: Template styles → Default styles → Custom styles
        final_styles = self._merge_styles(extracted_styles, styles)
        logger.info(f"[CONVERT] Step 3: Final styles merged: title_font={final_styles.get('title_font')}, body_font={final_styles.get('body_font')}, title_size={final_styles.get('title_size')}")
        
        try:
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                logger.info(f"[CONVERT] Creating output directory: {output_dir}")
                os.makedirs(output_dir, exist_ok=True)
            
            # Use base class parser
            logger.info("[CONVERT] Step 4: Parsing HTML elements")
            parser = HTMLElementParser()
            parser.feed(html)
            elements = parser.get_elements()
            logger.info(f"[CONVERT] Parsed {len(elements)} elements")
            
            # Log element types for debugging
            element_types = {}
            for elem in elements:
                elem_type = elem.get("type", "unknown")
                element_types[elem_type] = element_types.get(elem_type, 0) + 1
            logger.info(f"[CONVERT] Element breakdown: {element_types}")
            
            # Create Word document
            logger.info("[CONVERT] Step 5: Creating Word document")
            if self._docx_available:
                logger.info("[CONVERT] Using python-docx for document creation")
                result = self._create_docx_document(elements, output_path, final_styles)
            else:
                logger.warning("[CONVERT] python-docx not available, using fallback")
                result = self._create_fallback_document(elements, output_path, final_styles)
            
            if result.success:
                logger.info(f"[CONVERT] ========== Conversion SUCCESS ==========")
                logger.info(f"[CONVERT] Output: {result.output_path}")
                logger.info(f"[CONVERT] Size: {result.file_size} bytes")
                logger.info(f"[CONVERT] Pages estimate: {result.pages_estimate}")
            else:
                logger.error(f"[CONVERT] ========== Conversion FAILED ==========")
                logger.error(f"[CONVERT] Error: {result.error}")
                logger.error(f"[CONVERT] Error code: {result.error_code}")
            
            return result
                
        except (OSError, IOError) as e:
            logger.error(f"File operation failed: {e}")
            return ConversionResult(
                success=False,
                error=f"File operation failed: {e}",
                error_code="FILE_ERROR"
            )
        except ValueError as e:
            logger.error(f"Invalid data: {e}")
            return ConversionResult(
                success=False,
                error=f"Invalid data: {e}",
                error_code="DATA_ERROR"
            )
        except RuntimeError as e:
            logger.error(f"Processing error: {e}")
            return ConversionResult(
                success=False,
                error=f"Processing error: {e}",
                error_code="PROCESSING_ERROR"
            )
    
    def _sanitize_html(self, html: str) -> str:
        """
        Sanitize HTML content, remove content that shouldn't appear in Word
        
        1. Remove <style> tags and their content
        2. Remove residual template tags {% %} and {{ }}
        3. Remove other non-content tags
        4. **New**: Convert Markdown syntax to HTML tags
        
        Args:
            html: Original HTML content
            
        Returns:
            Sanitized HTML
        """
        import re
        
        # Remove <style> tags and content
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove <script> tags and content
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove <head> tags and content
        html = re.sub(r'<head[^>]*>.*?</head>', '', html, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove residual template tags
        html = re.sub(r'\{%.*?%\}', '', html, flags=re.DOTALL)
        html = re.sub(r'\{\{.*?\}\}', '', html, flags=re.DOTALL)
        
        # **New**: Convert Markdown syntax to HTML tags
        html = self._convert_markdown_to_html(html)
        
        # Clean extra whitespace
        html = re.sub(r'\n\s*\n\s*\n', '\n\n', html)
        
        return html.strip()
    
    def _convert_markdown_to_html(self, text: str) -> str:
        """
        Convert Markdown syntax to HTML tags
        
        Handled syntax:
        - **bold** → <strong>bold</strong>
        - *italic* → <em>italic</em>
        - `code` → <code>code</code>
        - [link](url) → <a href="url">link</a>
        - ~~strikethrough~~ → <del>strikethrough</del>
        
        Note: Only process text content between HTML tags, avoid breaking HTML structure
        
        Args:
            text: Text that may contain Markdown syntax
            
        Returns:
            Converted HTML
        """
        import re
        
        # Protect content in HTML tag attributes (avoid accidental conversion)
        # Strategy: Process in segments, only convert text outside tags
        
        # Simplified approach: Process Markdown syntax directly, but avoid processing inside HTML tags
        # Regex to detect if inside HTML tag: <...>
        
        # Split text into HTML tags and non-tag parts
        parts = re.split(r'(<[^>]+>)', text)
        
        result_parts = []
        for part in parts:
            # If it's an HTML tag, keep unchanged
            if part.startswith('<') and part.endswith('>'):
                result_parts.append(part)
            else:
                # Non-tag part, convert Markdown syntax
                converted = self._convert_markdown_inline(part)
                result_parts.append(converted)
        
        return ''.join(result_parts)
    
    def _convert_markdown_inline(self, text: str) -> str:
        """
        Convert inline Markdown syntax to HTML
        
        Processing order matters: process longer patterns first to avoid partial matches
        
        Args:
            text: Plain text content
            
        Returns:
            Converted HTML
        """
        import re
        
        # 1. Strikethrough ~~text~~
        text = re.sub(r'~~([^~]+)~~', r'<del>\1</del>', text)
        
        # 2. Bold **text** (must process before italic to avoid * being matched)
        text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
        
        # 3. Italic *text* or _text_
        text = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', text)
        text = re.sub(r'_([^_]+)_', r'<em>\1</em>', text)
        
        # 4. Inline code `code`
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        
        # 5. Link [text](url)
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
        
        return text
    
    def _is_safe_path(self, path: str, base_dir: Optional[str] = None) -> bool:
        """
        Check if path is safe
        
        Args:
            path: File path
            base_dir: Allowed base directory (optional, defaults to no restriction)
            
        Returns:
            Whether safe
        """
        try:
            resolved_path = Path(path).resolve()
            
            # Check path traversal ('..' as explicit check)
            if '..' in path:
                return False
            
            # If base directory specified, check if path is within allowed directory
            if base_dir is not None:
                base_path = Path(base_dir).resolve()
                
                # Check if absolute path is within allowed directory
                try:
                    if not resolved_path.is_relative_to(base_path):
                        return False
                except AttributeError:
                    # Fallback for Python < 3.9
                    try:
                        resolved_path.relative_to(base_path)
                    except ValueError:
                        return False
            
            # Check if path is absolute or valid relative path
            # Forbid access to system sensitive directories (exact match)
            dangerous_paths = [
                Path('/etc'), Path('/sys'), Path('/proc'), Path('/root'),
                Path('C:/Windows'), Path('C:/Windows/System32'),
                Path('D:/Windows'),  # Prevent other drive letters
            ]
            for dangerous in dangerous_paths:
                try:
                    if resolved_path.is_relative_to(dangerous):
                        return False
                except (AttributeError, ValueError):
                    # Fallback: Check path prefix
                    if str(resolved_path).startswith(str(dangerous)):
                        return False
            
            return True
            
        except (OSError, ValueError):
            return False
    
    def _atomic_save(self, doc: Any, output_path: str) -> None:
        """
        Atomically save document
        
        Use temp file + shutil.move to ensure atomic write.
        
        Args:
            doc: Document object
            output_path: Final output path
        """
        import shutil
        
        logger.info(f"[CONVERT] Atomic save starting, target: {output_path}")
        
        # Create temp file (don't use file descriptor)
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        temp_path = tempfile.mktemp(suffix='.docx', dir=output_dir or '.')
        logger.info(f"[CONVERT] Temp file path: {temp_path}")
        
        try:
            logger.info(f"[CONVERT] Saving document to temp file...")
            doc.save(temp_path)
            temp_size = os.path.getsize(temp_path) if os.path.exists(temp_path) else 0
            logger.info(f"[CONVERT] Temp file saved, size: {temp_size} bytes")
            
            logger.info(f"[CONVERT] Moving temp file to final location...")
            shutil.move(temp_path, output_path)
            logger.info(f"[CONVERT] Atomic save complete: {output_path}")
        except Exception as e:
            logger.error(f"[CONVERT] Atomic save FAILED: {e}")
            # Clean up temp file
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    logger.info(f"[CONVERT] Cleaned up temp file: {temp_path}")
            except OSError as cleanup_err:
                logger.warning(f"[CONVERT] Failed to cleanup temp file: {cleanup_err}")
            raise
    
    def _create_docx_document(
        self,
        elements: List[Dict[str, Any]],
        output_path: str,
        styles: Dict[str, Any]
    ) -> ConversionResult:
        """
        Create Word document using python-docx
        
        Args:
            elements: Parsed element list
            output_path: Output path
            styles: Style configuration
            
        Returns:
            ConversionResult
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set default font to Microsoft YaHei (unified for Chinese and English)
        style = doc.styles['Normal']
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(styles.get("body_size", 12))
        style._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
        # Apply body color (from CSS body color)
        body_color = styles.get("body_color", None)
        if body_color:
            rgb = self._parse_color_to_rgb(body_color)
            if rgb:
                font.color.rgb = RGBColor(*rgb)
        # Set paragraph spacing
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        
        paragraph_count = 0
        current_section_info = {}  # Current section info
        in_cover_area = True  # Whether in cover/TOC area (before first section)
        current_div_class = ""  # Current div class (cover-page/toc etc.)
        cover_title_added = False  # Whether cover title has been added
        
        for element in elements:
            elem_type = element.get("type", "")
            
            # Handle div markers: identify cover/TOC areas
            if elem_type == "div_start":
                current_div_class = element.get("class", "")
                # When cover starts, add top whitespace (centered slightly above)
                if current_div_class == "cover-page":
                    # 3 empty paragraphs ≈ 75px, make title at top 1/3 of page (centered slightly above)
                    for _ in range(3):
                        doc.add_paragraph("")
                continue
            elif elem_type == "div_end":
                # div end: add page break
                if current_div_class == "cover-page":
                    # Add page break after cover
                    doc.add_page_break()
                elif current_div_class == "toc":
                    # Add page break after TOC
                    doc.add_page_break()
                current_div_class = ""
                continue
            
            # Handle section markers: add page break before entering body area
            if elem_type == "section_start":
                if in_cover_area:
                    in_cover_area = False
                current_section_info = {
                    "id": element.get("id", ""),
                    "class": element.get("class", ""),
                    "data_type": element.get("data_type", "")
                }
                continue
            elif elem_type == "section_end":
                current_section_info = {}
                continue
            
            if elem_type == "heading":
                level = element.get("level", 1)
                text = element.get("text", "")
                css_class = element.get("style", "") or element.get("class", "")
                
                if text:
                    # Cover title special handling
                    if current_div_class == "cover-page" and level == 1 and not cover_title_added:
                        # Cover main title: large, centered, dark color
                        title_para = doc.add_paragraph()
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        title_run = title_para.add_run(text)
                        title_run.font.size = Pt(styles.get("title_size", 28))
                        title_run.font.name = "Microsoft YaHei"
                        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
                        title_run.font.bold = True
                        # Apply title color
                        title_color = styles.get("title_color", "#1C2B36")
                        rgb = self._parse_color_to_rgb(title_color)
                        if rgb:
                            title_run.font.color.rgb = RGBColor(*rgb)
                        cover_title_added = True
                        paragraph_count += 1
                    # TOC title special handling
                    elif current_div_class == "toc" and level == 2:
                        # TOC title: centered, large, letter spacing
                        toc_title = doc.add_paragraph()
                        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        toc_run = toc_title.add_run(text)
                        toc_run.font.size = Pt(18)
                        toc_run.font.name = "Microsoft YaHei"
                        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
                        toc_run.font.bold = True
                        rgb = self._parse_color_to_rgb(styles.get("title_color", "#1C2B36"))
                        if rgb:
                            toc_run.font.color.rgb = RGBColor(*rgb)
                        # Add empty line after TOC title
                        doc.add_paragraph("")
                        paragraph_count += 1
                    else:
                        # Normal heading
                        heading = doc.add_heading(text, level=level)
                        self._apply_heading_style(heading, level, styles, css_class)
                        paragraph_count += 1
            
            elif elem_type == "paragraph":
                text = element.get("text", "")
                para_class = element.get("class", "")
                
                if text:
                    para = doc.add_paragraph()
                    
                    # Cover area special handling
                    if current_div_class == "cover-page":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Subtitle style
                        if para_class == "subtitle":
                            run = para.add_run(text)
                            run.font.size = Pt(14)
                            run.font.name = "Microsoft YaHei"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
                            # Subtitle color
                            rgb = self._parse_color_to_rgb("#2C4A5A")
                            if rgb:
                                run.font.color.rgb = RGBColor(*rgb)
                        # Metadata style (author, date)
                        elif para_class == "meta":
                            run = para.add_run(text)
                            run.font.size = Pt(11)
                            run.font.name = "Microsoft YaHei"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
                            # Metadata color (gray)
                            rgb = self._parse_color_to_rgb("#888888")
                            if rgb:
                                run.font.color.rgb = RGBColor(*rgb)
                        else:
                            self._add_formatted_run(para, text, element.get("formats", []))
                    
                    # TOC area special handling
                    elif current_div_class == "toc":
                        # All paragraphs in TOC are TOC items (template rendering ensures this)
                        run = para.add_run(text)
                        run.font.size = Pt(12)
                        run.font.name = "Microsoft YaHei"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
                        # TOC item bottom dashed border effect (simulated with dots)
                        # python-docx doesn't support dashed borders, skip for now
                    
                    # Other areas
                    else:
                        self._add_formatted_run(para, text, element.get("formats", []))
                    
                    paragraph_count += 1
            
            elif elem_type == "list_item":
                text = element.get("text", "")
                list_type = element.get("list_type", "ul")
                
                if text:
                    if list_type == 'ol':
                        para = doc.add_paragraph(text, style='List Number')
                    else:
                        para = doc.add_paragraph(text, style='List Bullet')
                    paragraph_count += 1
            
            elif elem_type == "table":
                # Support both combined data and separate headers+rows formats
                data = element.get("data", [])
                if not data and element.get("headers") and element.get("rows"):
                    data = [element["headers"]] + element["rows"]
                
                if data:
                    # Calculate actual column count considering colspan
                    cols = 0
                    for row_data in data:
                        row_cols = 0
                        for cell_data in row_data:
                            if isinstance(cell_data, dict):
                                row_cols += cell_data.get("colspan", 1)
                            else:
                                row_cols += 1
                        cols = max(cols, row_cols)
                    
                    rows = len(data)
                    
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        try:
                            from docx.oxml.ns import qn
                            from docx.oxml import OxmlElement
                            from docx.shared import Pt, RGBColor
                            has_style_support = True
                        except ImportError:
                            has_style_support = False
                        
                        # Track absorbed cells to avoid overwriting merged content
                        absorbed = set()
                        
                        for i, row_data in enumerate(data):
                            actual_j = 0  # Actual column position (skip absorbed)
                            for j, cell_data in enumerate(row_data):
                                # Skip absorbed column positions
                                while actual_j < cols and (i, actual_j) in absorbed:
                                    actual_j += 1
                                if actual_j >= cols:
                                    break
                                
                                cell = table.rows[i].cells[actual_j]
                                
                                if isinstance(cell_data, dict):
                                    cell_text = str(cell_data.get("text", ""))
                                    colspan = cell_data.get("colspan", 1)
                                    rowspan = cell_data.get("rowspan", 1)
                                    is_header = cell_data.get("tag") == "th"
                                    
                                    # Merge cells
                                    if colspan > 1 or rowspan > 1:
                                        end_row = min(i + rowspan - 1, rows - 1)
                                        end_col = min(actual_j + colspan - 1, cols - 1)
                                        if end_row > i or end_col > actual_j:
                                            cell.merge(table.rows[end_row].cells[end_col])
                                            # Mark absorbed cell positions
                                            for ri in range(i, end_row + 1):
                                                for ci in range(actual_j, end_col + 1):
                                                    if ri != i or ci != actual_j:
                                                        absorbed.add((ri, ci))
                                            cell = table.rows[i].cells[actual_j]
                                else:
                                    cell_text = str(cell_data) if cell_data is not None else ""
                                    is_header = (i == 0)
                                
                                # Set text
                                cell.text = cell_text
                                
                                # Header style (dark background + white text + bold)
                                if is_header or i == 0:
                                    if has_style_support:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.bold = True
                                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                                run.font.size = Pt(10)
                                        shading = OxmlElement('w:shd')
                                        shading.set(qn('w:fill'), '1C2B36')
                                        cell._tc.get_or_add_tcPr().append(shading)
                                
                                actual_j += 1
                        
                        paragraph_count += 1
            
            # Image
            elif elem_type == "image":
                img_path = element.get("src", "")
                if img_path and os.path.exists(img_path):
                    try:
                        from docx.shared import Inches
                        # A4 body width ≈ 15cm ≈ 5.9in, after margins about 5.5in available
                        doc.add_picture(img_path, width=Inches(5.0))
                        # Add empty line after image last paragraph
                        last_para = doc.add_paragraph()
                        last_run = last_para.add_run("")
                        last_run.font.size = Pt(6)
                        paragraph_count += 1
                    except Exception as e:
                        logger.warning(f"Failed to add image {img_path}: {e}")
        
        # Atomic save
        logger.info(f"[CONVERT] Saving document, total paragraphs: {paragraph_count}")
        self._atomic_save(doc, output_path)
        
        # Get file size
        file_size = os.path.getsize(output_path)
        
        # Estimate page count
        pages_estimate = max(1, paragraph_count // 20)
        
        logger.info(f"[CONVERT] Word document created successfully:")
        logger.info(f"[CONVERT]   Path: {output_path}")
        logger.info(f"[CONVERT]   Size: {file_size} bytes ({file_size/1024:.1f} KB)")
        logger.info(f"[CONVERT]   Paragraphs: {paragraph_count}")
        logger.info(f"[CONVERT]   Pages estimate: {pages_estimate}")
        
        return ConversionResult(
            success=True,
            output_path=output_path,
            file_size=file_size,
            pages_estimate=pages_estimate
        )
    
    def _apply_heading_style(self, heading, level: int, styles: Dict[str, Any], css_class: str = "") -> None:
        """
        Apply CSS styles to heading
        
        Args:
            heading: Heading object
            level: Heading level
            styles: Style configuration
            css_class: CSS class name (to distinguish cover title/chapter title/section title)
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        
        # Decide which size to use based on CSS class name
        if css_class == "chapter-title":
            size_key = f"h{level}_size"
        elif css_class == "section-title":
            size_key = "h2_size"
        elif css_class == "subsection-title":
            size_key = "h3_size"
        elif level == 1 and not css_class:
            # Cover main title (<h1> without class) → use title_size (28pt)
            size_key = "title_size"
        else:
            size_key = f"h{level}_size"
        
        size = styles.get(size_key, styles.get("title_size", 24))
        color = styles.get("title_color", None)
        font_name = "Microsoft YaHei"  # Unified use of Microsoft YaHei
        
        for run in heading.runs:
            run.font.size = Pt(size)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if color:
                rgb = self._parse_color_to_rgb(color)
                if rgb:
                    run.font.color.rgb = RGBColor(*rgb)
    
    def _add_formatted_run(self, paragraph, text: str, formats: List[str]) -> None:
        """
        Add formatted text to paragraph
        
        Parse HTML inline tags in text (converted from Markdown)
        <strong>bold</strong> → bold
        <em>italic</em> → italic
        <code>code</code> → code
        
        Args:
            paragraph: Paragraph object
            text: Text that may contain HTML tags
            formats: Current format stack (passed from parser)
        """
        import re
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # If text still has HTML tags (from Markdown conversion), parse them
        # Use regex to extract tags and text
        pattern = r'<(strong|em|code|del|a[^>]*)>([^<]+)</(strong|em|code|del|a)>'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            # Add plain text before tag
            if match.start() > last_end:
                plain_text = text[last_end:match.start()]
                if plain_text:
                    run = paragraph.add_run(plain_text)
            
            # Add formatted text
            tag = match.group(1)
            content = match.group(2)
            
            run = paragraph.add_run(content)
            
            # Apply format
            if tag == 'strong':
                run.bold = True
            elif tag == 'em':
                run.italic = True
            elif tag == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif tag.startswith('a:'):
                # Link: extract URL
                pass  # Word doesn't support links, just display text
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            remaining = text[last_end:]
            # Clean other HTML tags
            remaining = re.sub(r'</?(strong|em|code|del|a[^>]*)>', '', remaining)
            if remaining:
                paragraph.add_run(remaining)
        
        # If no tags matched, add text directly
        if last_end == 0 and text:
            # Clean HTML tags
            clean_text = re.sub(r'</?(strong|em|code|del|a[^>]*)>', '', text)
            paragraph.add_run(clean_text)
    
    def _parse_color_to_rgb(self, color_str: str):
        """
        Convert color string to RGB tuple
        
        Args:
            color_str: Color string (e.g., "#1A2744" or "rgb(26, 39, 68)")
            
        Returns:
            (r, g, b) tuple or None
        """
        import re
        
        if not color_str:
            return None
        
        # Hex color
        if color_str.startswith('#'):
            hex_color = color_str[1:]
            if len(hex_color) == 6:
                return (
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16)
                )
        
        # RGB color
        match = re.match(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
        if match:
            return (int(match.group(1)), int(match.group(2)), int(match.group(3)))
        
        return None
    
    def _create_fallback_document(
        self,
        elements: List[Dict[str, Any]],
        output_path: str,
        styles: Dict[str, Any]
    ) -> ConversionResult:
        """
        Create fallback format document (when python-docx is not available)
        
        Args:
            elements: Parsed element list
            output_path: Output path
            styles: Style configuration
            
        Returns:
            ConversionResult
        """
        # Create plain text version
        text_lines = []
        
        for element in elements:
            elem_type = element.get("type", "")
            
            if elem_type == "heading":
                level = element.get("level", 1)
                text = element.get("text") or ""
                prefix = "#" * level
                text_lines.append(f"\n{prefix} {text}\n")
            
            elif elem_type == "paragraph":
                text = element.get("text") or ""
                text_lines.append(text)
            
            elif elem_type == "list_item":
                text = element.get("text") or ""
                list_type = element.get("list_type", "ul")
                if list_type == 'ol':
                    text_lines.append(f"1. {text}")
                else:
                    text_lines.append(f"- {text}")
            
            elif elem_type == "table":
                data = element.get("data", [])
                if not data and element.get("headers") and element.get("rows"):
                    data = [element["headers"]] + element["rows"]
                for row in data:
                    safe_row = [str(cell) if cell is not None else "" for cell in row]
                    text_lines.append(" | ".join(safe_row))
        
        content = "\n".join(text_lines)
        
        # Change to .txt extension
        txt_path = output_path.rsplit('.', 1)[0] + '.txt'
        
        # Atomic write
        fd, temp_path = tempfile.mkstemp(suffix='.txt')
        try:
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(content)
            os.replace(temp_path, txt_path)
        except Exception:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
            raise
        
        file_size = os.path.getsize(txt_path)
        
        logger.warning(f"python-docx not available, created text file: {txt_path}")
        
        return ConversionResult(
            success=True,
            output_path=txt_path,
            file_size=file_size,
            pages_estimate=len(text_lines) // 30,
            error="python-docx not available, created text file instead"
        )


# Export
__all__ = ["HTMLToWordConverter", "ConversionResult"]
