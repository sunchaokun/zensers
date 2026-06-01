# -*- coding: utf-8 -*-
"""
AI+ Manufacturing Deep Research Report - Professional Publishing Grade Generator
================================================================================

Strategy:
1. Directly read all markdown source files (full content)
2. Apply all professional styles from template_generator.js
3. Generate publishing-grade Word documents
"""

import re
import os
from docx import Document
import docx
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# ============================================================
# Publishing-grade professional color system (from template_generator.js)
# ============================================================
COLORS = {
    'primaryDark': RGBColor(0x1A, 0x27, 0x44),   # Dark navy blue
    'primary': RGBColor(0x2C, 0x3E, 0x50),        # Classic blue-gray
    'primaryLight': RGBColor(0x34, 0x49, 0x5E),   # Light blue-gray
    'accent': RGBColor(0xC9, 0xA2, 0x27),         # Amber gold
    'accentLight': RGBColor(0xD4, 0xAF, 0x37),   # Bright gold
    'accentDark': RGBColor(0xA6, 0x8B, 0x1F),   # Deep gold
    'success': RGBColor(0x27, 0xAE, 0x60),        # Professional green
    'warning': RGBColor(0xE6, 0x7E, 0x22),        # Warning orange
    'danger': RGBColor(0xC0, 0x39, 0x2B),         # Risk red
    'info': RGBColor(0x34, 0x98, 0xDB),           # Information blue
    'white': RGBColor(0xFF, 0xFF, 0xFF),
    'offWhite': RGBColor(0xF8, 0xF9, 0xFA),
    'lightGray': RGBColor(0xEC, 0xF0, 0xF1),
    'mediumGray': RGBColor(0x95, 0xA5, 0xA6),
    'darkGray': RGBColor(0x7F, 0x8C, 0x8D),
    'textDark': RGBColor(0x2C, 0x3E, 0x50),
    'text': RGBColor(0x33, 0x33, 0x33),
    'textLight': RGBColor(0x66, 0x66, 0x66),
    'tableHeader': RGBColor(0x1A, 0x27, 0x44),
    'tableBorder': RGBColor(0xBD, 0xC3, 0xC7),
    'tableAlt': RGBColor(0xF8, 0xF9, 0xFA),
    # Hex string versions (for shading)
    'primaryDarkHex': '1A2744',
    'primaryHex': '2C3E50',
    'primaryLightHex': '34495E',
    'accentHex': 'C9A227',
    'successHex': '27AE60',
    'warningHex': 'E67E22',
    'dangerHex': 'C0392B',
    'whiteHex': 'FFFFFF',
    'offWhiteHex': 'F8F9FA',
    'lightGrayHex': 'ECF0F1',
    'mediumGrayHex': '95A5A6',
    'darkGrayHex': '7F8C8D',
    'textDarkHex': '2C3E50',
    'textHex': '333333',
    'textLightHex': '666666',
    'tableHeaderHex': '1A2744',
    'tableBorderHex': 'BDC3C7',
    'tableAltHex': 'F8F9FA',
    'cardBgHex': 'FDF8E8',  # Key finding card background
    'caseCardBgHex': 'F8F9FA',  # Case card background
}

# Fonts
FONTS = {
    'title': 'Georgia',
    'body': 'Calibri',
    'accent': 'Arial',
    'mono': 'Courier New',
}

# ============================================================
# Utility functions
# ============================================================
def rgb_to_hex(rgb):
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), edge_data.get('val', 'single'))
            tag.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            tag.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_paragraph_border_bottom(paragraph, color_hex, size=6, space=8):
    """Set paragraph bottom border"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_border_left(paragraph, color_hex, size=24, space=4):
    """Set paragraph left border (gold accent line)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def set_paragraph_shading(paragraph, fill_color):
    """Set paragraph background"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)

def add_run_font(run, font_name):
    """Set run font"""
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

# ============================================================
# Paragraph style creation functions
# ============================================================
def create_heading1(doc, text):
    """Heading 1: Dark navy blue, Georgia, 18pt - Template standard"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '400')
    spacing.set(qn('w:after'), '200')
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLORS['primaryDark']
    run.font.name = FONTS['title']
    add_run_font(run, FONTS['title'])
    return p

def create_heading2(doc, text):
    """Heading 2: Classic blue-gray, Georgia"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '360')
    spacing.set(qn('w:after'), '160')
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)  # 28 half-points = 14pt
    run.font.color.rgb = COLORS['primary']
    run.font.name = FONTS['title']
    add_run_font(run, FONTS['title'])
    return p

def create_heading3(doc, text):
    """Heading 3: Light blue-gray, Calibri"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)  # 24 half-points = 12pt
    run.font.color.rgb = COLORS['primaryLight']
    run.font.name = FONTS['body']
    add_run_font(run, FONTS['body'])
    return p

def create_heading4(doc, text):
    """Heading 4: Blue-gray, Calibri, no top/bottom margin"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '180')
    spacing.set(qn('w:after'), '100')
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)  # 22 half-points = 11pt
    run.font.color.rgb = COLORS['textDark']
    run.font.name = FONTS['body']
    add_run_font(run, FONTS['body'])
    return p

def create_paragraph(doc, text, options=None):
    """Body paragraph: Calibri, justified alignment"""
    if options is None:
        options = {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), str(options.get('spacingAfter', 120)))
    pPr.append(spacing)
    run = p.add_run(text)
    run.font.size = Pt(10)  # Template standard: 10pt
    run.font.color.rgb = COLORS['text']
    run.font.name = FONTS['body']
    run.bold = options.get('bold', False)
    run.italic = options.get('italic', False)
    add_run_font(run, FONTS['body'])
    if options.get('color'):
        run.font.color.rgb = options['color']
    return p

def create_bullet(doc, text, level=0):
    """Bullet point"""
    indent_left = 560 if level == 0 else 840
    bullet_char = '\u2022' if level == 0 else '\u25E6'
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '80')
    pPr.append(spacing)
    # Indentation
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_left))
    ind.set(qn('w:hanging'), '280')
    pPr.append(ind)
    run = p.add_run(f'{bullet_char}  {text}')
    run.font.size = Pt(10)  # Template standard: 10pt
    run.font.color.rgb = COLORS['text']
    run.font.name = FONTS['body']
    add_run_font(run, FONTS['body'])
    return p

def create_numbered(doc, text, number):
    """Numbered paragraph"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '80')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '560')
    ind.set(qn('w:hanging'), '280')
    pPr.append(ind)
    # Bold number prefix
    run1 = p.add_run(f'{number}. ')
    run1.bold = True
    run1.font.size = Pt(10)  # Template standard: 10pt
    run1.font.color.rgb = COLORS['primary']
    run1.font.name = FONTS['body']
    add_run_font(run1, FONTS['body'])
    # Normal content
    run2 = p.add_run(text)
    run2.font.size = Pt(10)  # Template standard: 10pt
    run2.font.color.rgb = COLORS['text']
    run2.font.name = FONTS['body']
    add_run_font(run2, FONTS['body'])
    return p

def create_separator(doc, color_hex=None, size=3):
    """Separator line"""
    if color_hex is None:
        color_hex = COLORS['lightGrayHex']
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '200')
    pPr.append(spacing)
    set_paragraph_border_bottom(p, color_hex, size, 4)
    return p

def create_space(doc, points=120):
    """Empty line"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), str(points))
    pPr.append(spacing)
    return p

def create_page_break(doc):
    """Page break"""
    p = doc.add_paragraph()
    run = p.add_run()
    # Add page break directly in run
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return p

# ============================================================
# Table components
# ============================================================
def create_professional_table(doc, headers, rows, widths, options=None):
    """Professional table: Dark blue header, alternating row colors"""
    if options is None:
        options = {}
    header_color = options.get('headerColor', COLORS['tableHeaderHex'])
    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for i, w in enumerate(widths):
        table.columns[i].width = Twips(w)

    # Header row
    header_row = table.add_row()
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, header_color)
        set_cell_border(cell,
            top={'val': 'single', 'sz': 4, 'color': COLORS['tableBorderHex']},
            bottom={'val': 'single', 'sz': 4, 'color': COLORS['tableBorderHex']},
            left={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
            right={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
        )
        # Cell padding
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom', 'left', 'right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '120' if side in ['left', 'right'] else '100')
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        # Vertical center
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
            # Content
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(options.get('headerSize', 9.5))
        run.font.color.rgb = COLORS['white']
        run.font.name = FONTS['body']
        add_run_font(run, FONTS['body'])

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        is_alt = row_idx % 2 == 1
        bg = COLORS['tableAltHex'] if is_alt else COLORS['whiteHex']
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_shading(cell, bg)
            set_cell_border(cell,
                top={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                bottom={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                left={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                right={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
            )
            # Padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '120' if side in ['left', 'right'] else '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            # Content
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLORS['text']
            run.font.name = FONTS['body']
            run.bold = options.get('boldFirstCol', False) and col_idx == 0
            add_run_font(run, FONTS['body'])

    return table

def create_simple_table(doc, rows, widths, options=None):
    """Simple table (no header style distinction)"""
    if options is None:
        options = {}
    table = doc.add_table(rows=0, cols=len(rows[0]) if rows else 0)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for col_idx, w in enumerate(widths):
        table.columns[col_idx].width = Twips(w)

    for row_idx, row_data in enumerate(row_data for row_data in rows):
        row = table.add_row()
        is_header = options.get('headerRow', False) and row_idx == 0
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            if is_header:
                set_cell_shading(cell, options.get('headerColor', COLORS['tableHeaderHex']))
            else:
                is_alt = row_idx % 2 == 1
                set_cell_shading(cell, COLORS['tableAltHex'] if is_alt else COLORS['whiteHex'])
            set_cell_border(cell,
                top={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                bottom={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                left={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
                right={'val': 'single', 'sz': 2, 'color': COLORS['tableBorderHex']},
            )
            # Padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '120' if side in ['left', 'right'] else '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.bold = is_header
            if is_header:
                run.font.color.rgb = COLORS['white']
            else:
                run.font.color.rgb = COLORS['text']
            run.font.name = FONTS['body']
            add_run_font(run, FONTS['body'])

    return table

# ============================================================
# Chapter cover
# ============================================================
def create_chapter_cover(doc, chapter_num, title, subtitle_en):
    """Chapter cover"""
    elements = []

    # Top margin space
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '1440')
    pPr.append(spacing)

    # Gold separator line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run2.font.size = Pt(10)
    run2.font.color.rgb = COLORS['accent']
    run2.font.name = FONTS['body']
    add_run_font(run2, FONTS['body'])

    # Chapter number
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr3 = p3._p.get_or_add_pPr()
    spacing3 = OxmlElement('w:spacing')
    spacing3.set(qn('w:before'), '200')
    pPr3.append(spacing3)
    run3 = p3.add_run(f'Chapter {chapter_num}')
    run3.bold = True
    run3.font.size = Pt(22)
    run3.font.color.rgb = COLORS['primaryDark']
    run3.font.name = FONTS['title']
    add_run_font(run3, FONTS['title'])

    # Chinese title
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr4 = p4._p.get_or_add_pPr()
    spacing4 = OxmlElement('w:spacing')
    spacing4.set(qn('w:before'), '120')
    pPr4.append(spacing4)
    run4 = p4.add_run(title)
    run4.bold = True
    run4.font.size = Pt(28)
    run4.font.color.rgb = COLORS['primaryDark']
    run4.font.name = FONTS['title']
    add_run_font(run4, FONTS['title'])

    # English subtitle
    if subtitle_en:
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr5 = p5._p.get_or_add_pPr()
        spacing5 = OxmlElement('w:spacing')
        spacing5.set(qn('w:before'), '80')
        pPr5.append(spacing5)
        run5 = p5.add_run(subtitle_en)
        run5.font.size = Pt(12)
        run5.font.color.rgb = COLORS['mediumGray']
        run5.font.name = FONTS['body']
        add_run_font(run5, FONTS['body'])

    # Gold separator line
    p6 = doc.add_paragraph()

    return elements


# ============================================================
# Key finding card (emphasized area with gold left border)
# ============================================================
def create_key_finding_card(doc, number, title, data_points, insights, uncertainty):
    """Key finding card"""
    # Title (with gold left border)
    p1 = doc.add_paragraph()
    pPr1 = p1._p.get_or_add_pPr()
    spacing1 = OxmlElement('w:spacing')
    spacing1.set(qn('w:before'), '280')
    pPr1.append(spacing1)
    set_paragraph_border_left(p1, COLORS['accentHex'], 24, 4)
    set_paragraph_shading(p1, COLORS['cardBgHex'])
    ind1 = OxmlElement('w:ind')
    ind1.set(qn('w:left'), '200')
    pPr1.append(ind1)
    run1 = p1.add_run(f'{number}. {title}')
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = COLORS['primaryDark']
    run1.font.name = FONTS['title']
    add_run_font(run1, FONTS['title'])

    # Data support label
    p2 = doc.add_paragraph()
    pPr2 = p2._p.get_or_add_pPr()
    set_paragraph_shading(p2, COLORS['cardBgHex'])
    ind2 = OxmlElement('w:ind')
    ind2.set(qn('w:left'), '200')
    pPr2.append(ind2)
    spacing2 = OxmlElement('w:spacing')
    spacing2.set(qn('w:before'), '80')
    spacing2.set(qn('w:after'), '60')
    pPr2.append(spacing2)
    run2 = p2.add_run('Data Support')
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = COLORS['accent']
    run2.font.name = FONTS['body']
    add_run_font(run2, FONTS['body'])

    # Data points
    for dp in data_points:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        set_paragraph_shading(p, COLORS['cardBgHex'])
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '400')
        pPr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        run = p.add_run(f'• {dp}')
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLORS['text']
        run.font.name = FONTS['body']
        add_run_font(run, FONTS['body'])

    # Key insight label
    p3 = doc.add_paragraph()
    pPr3 = p3._p.get_or_add_pPr()
    set_paragraph_shading(p3, COLORS['cardBgHex'])
    ind3 = OxmlElement('w:ind')
    ind3.set(qn('w:left'), '200')
    pPr3.append(ind3)
    spacing3 = OxmlElement('w:spacing')
    spacing3.set(qn('w:before'), '100')
    spacing3.set(qn('w:after'), '60')
    pPr3.append(spacing3)
    run3 = p3.add_run('Key Insight')
    run3.bold = True
    run3.font.size = Pt(9)
    run3.font.color.rgb = COLORS['primary']
    run3.font.name = FONTS['body']
    add_run_font(run3, FONTS['body'])

    # Insight content
    for insight in insights:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        set_paragraph_shading(p, COLORS['cardBgHex'])
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '400')
        pPr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)

        prefix = '[OK] '
        color = COLORS['success']
        text = insight
        if insight.startswith('[WARNING]'):
            prefix = '[!] '
            color = COLORS['warning']
            text = insight.replace('[WARNING] ', '')
        elif insight.startswith('[RISK]') or insight.startswith('[FAILED]'):
            prefix = '[X] '
            color = COLORS['danger']
            text = insight.replace('[RISK] ', '').replace('[FAILED] ', '')

        run = p.add_run(f'{prefix}{text}')
        run.font.size = Pt(9.5)
        run.font.color.rgb = color
        run.font.name = FONTS['body']
        add_run_font(run, FONTS['body'])

    # Uncertainty assessment
    p4 = doc.add_paragraph()
    pPr4 = p4._p.get_or_add_pPr()
    set_paragraph_shading(p4, COLORS['cardBgHex'])
    ind4 = OxmlElement('w:ind')
    ind4.set(qn('w:left'), '200')
    pPr4.append(ind4)
    spacing4 = OxmlElement('w:spacing')
    spacing4.set(qn('w:before'), '80')
    spacing4.set(qn('w:after'), '200')
    pPr4.append(spacing4)
    run4a = p4.add_run('Uncertainty Assessment: ')
    run4a.bold = True
    run4a.font.size = Pt(9)
    run4a.font.color.rgb = COLORS['darkGray']
    run4a.font.name = FONTS['body']
    add_run_font(run4a, FONTS['body'])
    run4b = p4.add_run(uncertainty)
    run4b.font.size = Pt(9)
    run4b.font.color.rgb = COLORS['primary']
    run4b.font.name = FONTS['body']
    add_run_font(run4b, FONTS['body'])

# ============================================================
# Case card
# ============================================================
def create_case_card(doc, title, sections):
    """Case card (with dark blue left border)"""
    # Title
    p1 = doc.add_paragraph()
    pPr1 = p1._p.get_or_add_pPr()
    spacing1 = OxmlElement('w:spacing')
    spacing1.set(qn('w:before'), '280')
    spacing1.set(qn('w:after'), '100')
    pPr1.append(spacing1)
    set_paragraph_border_left(p1, COLORS['primaryHex'], 18, 4)
    set_paragraph_shading(p1, COLORS['offWhiteHex'])
    ind1 = OxmlElement('w:ind')
    ind1.set(qn('w:left'), '150')
    pPr1.append(ind1)
    run1 = p1.add_run(title)
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = COLORS['primaryDark']
    run1.font.name = FONTS['title']
    add_run_font(run1, FONTS['title'])

    for section in sections:
        if section['type'] == 'label':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            set_paragraph_shading(p, COLORS['offWhiteHex'])
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '300')
            pPr.append(ind)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '120')
            spacing.set(qn('w:after'), '50')
            pPr.append(spacing)
            run = p.add_run(section['label'])
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLORS['primary']
            run.font.name = FONTS['body']
            add_run_font(run, FONTS['body'])
        elif section['type'] == 'bullet':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            set_paragraph_shading(p, COLORS['offWhiteHex'])
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '300')
            pPr.append(ind)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:after'), '40')
            pPr.append(spacing)
            run = p.add_run(f'• {section["content"]}')
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLORS['text']
            run.font.name = FONTS['body']
            add_run_font(run, FONTS['body'])

# ============================================================
# Cover
# ============================================================
def create_cover(doc, title, subtitle, date, version):
    """Report cover"""
    # Internal classification label
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr1 = p1._p.get_or_add_pPr()
    spacing1 = OxmlElement('w:spacing')
    spacing1.set(qn('w:before'), '1200')
    pPr1.append(spacing1)
    run1 = p1.add_run('[ Internal Research Edition ]')
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.color.rgb = COLORS['accent']
    run1.font.name = FONTS['body']
    add_run_font(run1, FONTS['body'])

    # Gold separator line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p2._p.get_or_add_pPr()
    spacing2 = OxmlElement('w:spacing')
    spacing2.set(qn('w:before'), '80')
    pPr2.append(spacing2)
    run2 = p2.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run2.font.size = Pt(10)
    run2.font.color.rgb = COLORS['accent']
    run2.font.name = FONTS['body']
    add_run_font(run2, FONTS['body'])

    # English label
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr3 = p3._p.get_or_add_pPr()
    spacing3 = OxmlElement('w:spacing')
    spacing3.set(qn('w:before'), '120')
    pPr3.append(spacing3)
    run3 = p3.add_run('DEEP INDUSTRY RESEARCH REPORT')
    run3.font.size = Pt(8)
    run3.font.color.rgb = COLORS['darkGray']
    run3.font.name = FONTS['body']
    add_run_font(run3, FONTS['body'])

    # Main title
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr4 = p4._p.get_or_add_pPr()
    spacing4 = OxmlElement('w:spacing')
    spacing4.set(qn('w:before'), '320')
    pPr4.append(spacing4)
    run4 = p4.add_run(title)
    run4.bold = True
    run4.font.size = Pt(26)
    run4.font.color.rgb = COLORS['primaryDark']
    run4.font.name = FONTS['title']
    add_run_font(run4, FONTS['title'])

    # Subtitle
    if subtitle:
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr5 = p5._p.get_or_add_pPr()
        spacing5 = OxmlElement('w:spacing')
        spacing5.set(qn('w:before'), '160')
        pPr5.append(spacing5)
        run5 = p5.add_run(subtitle)
        run5.font.size = Pt(13)
        run5.font.color.rgb = COLORS['primary']
        run5.font.name = FONTS['body']
        add_run_font(run5, FONTS['body'])

    # Gold separator line
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr6 = p6._p.get_or_add_pPr()
    spacing6 = OxmlElement('w:spacing')
    spacing6.set(qn('w:before'), '480')
    pPr6.append(spacing6)
    run6 = p6.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run6.font.size = Pt(10)
    run6.font.color.rgb = COLORS['accent']
    run6.font.name = FONTS['body']
    add_run_font(run6, FONTS['body'])

    # Empty line
    create_space(doc, 200)

    # Cover table
    cover_rows = [
        ['Research Type', 'Deep Industry Research'],
        ['Data Cutoff', date],
        ['Version', version],
        ['Confidence Label', 'Level 1 & 2 Data > 95%'],
    ]
    cover_widths = [2000, 7540]
    table = create_simple_table(doc, cover_rows, cover_widths)
    # Header style
    for row in table.rows:
        cell = row.cells[0]
        set_cell_shading(cell, COLORS['tableHeaderHex'])
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.color.rgb = COLORS['white']
            run.bold = True

    # Empty line
    create_space(doc, 600)

    # Footer info
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr7 = p7._p.get_or_add_pPr()
    spacing7 = OxmlElement('w:spacing')
    spacing7.set(qn('w:before'), '60')
    pPr7.append(spacing7)
    run7 = p7.add_run('AI+ Manufacturing Deep Research Report  |  Internal Research Material')
    run7.font.size = Pt(8)
    run7.font.color.rgb = COLORS['darkGray']
    run7.font.name = FONTS['body']
    add_run_font(run7, FONTS['body'])

# ============================================================
# Markdown Parser
# ============================================================
def parse_markdown_table(table_text):
    """Parse markdown table, returns (headers, rows)"""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return [], []

    headers = [h.strip() for h in lines[0].split('|')[1:-1]]
    rows = []
    for line in lines[2:]:  # Skip separator line
        if '|' in line and not set(line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')) == set():
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and any(cells):
                rows.append(cells)
    return headers, rows

def get_table_widths(headers, rows):
    """Estimate column widths based on content (total width ~9540 DXA)"""
    n_cols = len(headers)
    if n_cols == 0:
        return []
    total = 9540
    base = total // n_cols
    widths = [base] * n_cols
    # Adjust
    widths[0] = int(base * 0.6)
    remaining = total - widths[0]
    per_col = remaining // (n_cols - 1)
    for i in range(1, n_cols):
        widths[i] = per_col
    return widths

def markdown_to_docx(doc, content):
    """Convert markdown content to docx elements and add to document"""
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].strip()

        # Table processing
        if '|' in line and not line.startswith('#'):
            table_lines.append(line)
            # Predict if still in table
            next_is_table = (i + 1 < len(lines) and '|' in lines[i+1].strip())
            next_is_sep = (i + 1 < len(lines) and set(lines[i+1].strip().replace('|', '').replace('-', '').replace(':', '').replace(' ', '')) == set())
            if not next_is_table:
                # Parse table
                headers, rows = parse_markdown_table('\n'.join(table_lines))
                if headers and rows:
                    widths = get_table_widths(headers, rows)
                    create_professional_table(doc, headers, rows, widths)
                table_lines = []
                in_table = False
            i += 1
            continue
        elif in_table and not table_lines:
            in_table = False

        # Heading 1
        if line.startswith('# '):
            create_heading1(doc, line[2:].strip())
        # Heading 2
        elif line.startswith('## '):
            create_heading2(doc, line[3:].strip())
        # Heading 3
        elif line.startswith('### '):
            create_heading3(doc, line[4:].strip())
        # Heading 4
        elif line.startswith('#### '):
            create_heading4(doc, line[5:].strip())
        # Separator
        elif line.startswith('---') or line.startswith('***') or line.startswith('___'):
            create_separator(doc)
        # Page break
        elif line.startswith('<!-- pagebreak -->') or line == '{pagebreak}':
            create_page_break(doc)
        # List item
        elif line.startswith('- ') or line.startswith('* '):
            create_bullet(doc, line[2:].strip())
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            match = re.match(r'^(\d+)\. (.+)', line)
            if match:
                create_numbered(doc, match.group(2), int(match.group(1)))
        # Empty line
        elif not line:
            pass  # Skip empty lines
        # Plain text
        else:
            # Clean residual markdown
            text = line
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)       # Italic
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Link
            if text.strip():
                create_paragraph(doc, text)

        i += 1

# ============================================================
# Main function
# ============================================================
def build_report():
    base_dir = r'e:\douyin\创作\AI+制造业深度研究报告'

    # Read all markdown content
    md_files = [
        '2026-03-28_AI+制造业深度报告_第1-2章.md',
        '2026-03-28_AI+制造业深度报告_第3章.md',
        '2026-03-28_AI+制造业深度报告_第4章.md',
        '2026-03-28_AI+制造业深度报告_第5-8章.md',
        '2026-03-28_AI+制造业深度报告_附录1-4.md',
    ]

    # Read source file contents (as raw content)
    source_files = [
        '2026-03-28_AI+制造业深度报告_第1-2章.md',
        '2026-03-28_AI+制造业深度报告_第3章.md',
        '2026-03-28_AI+制造业深度报告_第4章.md',
        '2026-03-28_AI+制造业深度报告_第5-8章.md',
        '2026-03-28_AI+制造业深度报告_附录1-4.md',
    ]

    contents = []
    total_chars = 0
    for fname in source_files:
        fpath = os.path.join(base_dir, fname)
        if os.path.exists(fpath):
            with open(fpath, 'r', encoding='utf-8') as f:
                content = f.read()
                contents.append((fname, content))
                total_chars += len(content)
                print(f'  Read: {fname} ({len(content):,} chars)')
        else:
            print(f'  [WARN] File not found: {fname}')

    print(f'\nTotal content: {total_chars:,} chars from {len(contents)} files')

    # Create document
    doc = Document()

    # Set page margins (top/bottom 2.54cm, left/right 3.18cm)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Set default paragraph style
    doc.styles['Normal'].font.name = FONTS['body']
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

    # ========== Cover ==========
    create_cover(doc,
        title='AI+ Manufacturing Deep Research Report',
        subtitle='Strategic Opportunities and Implementation Paths in the Intelligent Transformation Wave',
        date='March 28, 2026',
        version='v1.0'
    )

    create_page_break(doc)

    # ========== Process files sequentially ==========
    for fname, content in contents:
        print(f'\nProcessing: {fname}')
        markdown_to_docx(doc, content)
        print(f'  Done: {len(content):,} chars converted')

    # ========== Report End ==========
    create_space(doc, 300)
    create_separator(doc, COLORS['accentHex'], 4)

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_end._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    pPr.append(spacing)
    run = p_end.add_run('— End of Report —')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS['primaryDark']
    run.font.name = FONTS['title']
    add_run_font(run, FONTS['title'])

    p_copy = doc.add_paragraph()
    p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p_copy._p.get_or_add_pPr()
    spacing2 = OxmlElement('w:spacing')
    spacing2.set(qn('w:before'), '100')
    pPr2.append(spacing2)
    run2 = p_copy.add_run('(c) 2026 AI Knowledge System  |  This report is for reference only and does not constitute investment advice')
    run2.font.size = Pt(8)
    run2.font.color.rgb = COLORS['darkGray']
    run2.font.name = FONTS['body']
    add_run_font(run2, FONTS['body'])

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr3 = p_date._p.get_or_add_pPr()
    spacing3 = OxmlElement('w:spacing')
    spacing3.set(qn('w:before'), '60')
    pPr3.append(spacing3)
    run3 = p_date.add_run('Data Updated: March 28, 2026  |  Next Update: June 2026 (Quarterly Update)')
    run3.font.size = Pt(7)
    run3.font.color.rgb = COLORS['mediumGray']
    run3.font.name = FONTS['body']
    add_run_font(run3, FONTS['body'])

    # ========== Save ==========
    output_path = os.path.join(base_dir, 'AI制造业深度研究报告_完整版.docx')
    doc.save(output_path)

    # Verify
    size = os.path.getsize(output_path)
    print(f'\n{"="*50}')
    print(f'Document saved: {output_path}')
    print(f'File size: {size/1024:.1f} KB')
    print(f'Content chars: {total_chars:,}')
    print(f'{"="*50}')

    return output_path

if __name__ == '__main__':
    build_report()
