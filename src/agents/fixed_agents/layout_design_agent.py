"""
Layout Design Agent
==================

Responsible for report formatting output, supporting multiple formats.

Responsibilities:
1. Convert Markdown reports to Word documents
2. Apply professional styles (fonts, colors, spacing)
3. Insert charts and tables
4. Generate table of contents, headers and footers

Input:
{
    "content": str,             # Markdown format report content
    "output_format": str,       # Output format: docx|pdf|html
    "style_config": dict,       # Style configuration (optional)
    "charts": list,             # Chart configuration (optional)
    "tables": list,             # Table configuration (optional)
}

Output:
{
    "success": bool,
    "file_path": str,           # Generated file path
    "file_size": int,           # File size (bytes)
    "format": str,              # Actual output format
    "pages": int,               # Page count estimate
}
"""

import os
from typing import Any, Dict, List, Optional
from datetime import datetime
from .base_fixed_agent import FixedAgent

# P0-3 fix: Move docx-related imports to module level
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class LayoutDesignAgent(FixedAgent):
    """Layout Design Agent.
    
    Responsible for converting report content to professionally formatted output documents.
    Supports Word, PDF, HTML and other formats.
    """
    
    agent_type = "layout_design"
    version = "1.0.0"
    capabilities = [
        "Word Document Generation",
        "PDF Export",
        "HTML Generation",
        "Style Application",
        "Chart Insertion",
        "Table of Contents Generation",
    ]
    
    # Default style configuration
    DEFAULT_STYLE = {
        "colors": {
            "primaryDark": "1A2744",   # Deep navy blue
            "primary": "2C3E50",        # Classic blue-gray
            "accent": "C9A227",         # Amber gold
            "success": "27AE60",        # Professional green
            "warning": "E67E22",        # Warning orange
        },
        "fonts": {
            "title": "Georgia",
            "body": "Calibri",
            "chinese": "SimSun",
        },
        "margins": {
            "top": 2.54,    # cm
            "bottom": 2.54,
            "left": 3.18,
            "right": 3.18,
        },
        "line_spacing": 1.5,
        "table_width": 100,
    }
    
    def validate_input(self, task_input: Dict[str, Any]) -> tuple[bool, str]:
        """Validate input parameters."""
        valid, error = super().validate_input(task_input)
        if not valid:
            return valid, error
        
        if "content" not in task_input:
            return False, "Missing required field 'content'"
        
        output_format = task_input.get("output_format", "docx")
        if output_format not in ["docx", "pdf", "html", "markdown"]:
            return False, f"Unsupported output format: {output_format}"
        
        return True, ""
    
    async def execute(self, task_input: Dict[str, Any]) -> Dict[str, Any]:
        """Execute layout design (async).
        
        Args:
            task_input: {
                "content": "# Report Title\n...",
                "output_format": "docx",
                "style_config": {...},
                "output_path": "/path/to/output.docx",
            }
        """
        content = task_input["content"]
        output_format = task_input.get("output_format", "docx")
        # P0-3 fix: Properly handle None value, use or operator
        style_config = task_input.get("style_config") or self.DEFAULT_STYLE
        output_path = task_input.get("output_path")
        
        # Publish start event
        await self.publish_event("layout_started", {"format": output_format})
        
        # Generate default path if output path not specified
        if not output_path:
            output_path = self._generate_output_path(output_format)
        
        # Select generation method based on format
        if output_format == "docx":
            result = self._generate_docx(content, output_path, style_config)
        elif output_format == "html":
            result = self._generate_html(content, output_path, style_config)
        elif output_format == "markdown":
            result = self._save_markdown(content, output_path)
        else:
            result = {"success": False, "error": f"Format not yet supported: {output_format}"}
        
        # Publish completion event
        await self.publish_event("layout_completed", {"output_path": output_path})
        
        return result
    
    def _generate_output_path(self, output_format: str) -> str:
        """Generate default output path."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.{output_format}"
        
        # Use storage_path or current directory
        if self.storage_path:
            return os.path.join(self.storage_path, filename)
        return filename
    
    def _generate_docx(
        self, 
        content: str, 
        output_path: str, 
        style_config: Dict
    ) -> Dict[str, Any]:
        """Generate Word document.
        
        Uses python-docx library to generate professionally formatted Word documents.
        """
        # P0-3 fix: Use module-level import
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx library not found, install: pip install python-docx"
            }
        
        # Create document
        doc = Document()
        
        # Set default font
        self._setup_docx_styles(doc, style_config)
        
        # Set page margins
        sections = doc.sections[0]
        margins = style_config.get("margins", self.DEFAULT_STYLE["margins"])
        sections.top_margin = Cm(margins["top"])
        sections.bottom_margin = Cm(margins["bottom"])
        sections.left_margin = Cm(margins["left"])
        sections.right_margin = Cm(margins["right"])
        
        # Parse Markdown and convert to Word
        self._markdown_to_docx(doc, content, style_config)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # Save document
        doc.save(output_path)
        
        # Get file information
        file_size = os.path.getsize(output_path)
        
        return {
            "success": True,
            "file_path": output_path,
            "file_size": file_size,
            "format": "docx",
            "pages": "Pending calculation",  # python-docx does not directly provide page count
        }
    
    def _setup_docx_styles(self, doc, style_config: Dict):
        """Set up Word document styles."""
        fonts = style_config.get("fonts", self.DEFAULT_STYLE["fonts"])
        
        # Set Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = fonts.get("body", "Calibri")
        font.size = Pt(11)
        
        # Create heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = fonts.get("title", "Georgia")
            heading_style.font.size = Pt([18, 14, 12][i-1])
            heading_style.font.bold = True
    
    def _markdown_to_docx(self, doc, content: str, style_config: Dict):
        """Convert Markdown to Word format.
        
        Simplified implementation supporting basic Markdown syntax.
        """
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            
            # Horizontal rule
            elif line == '---':
                doc.add_paragraph('_' * 50)
            
            # List items
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. '):
                doc.add_paragraph(line[3:], style='List Number')
            
            # Bold and italic
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            
            # Regular paragraph
            else:
                # Handle inline formatting
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text."""
        # Simplified processing, can support more complex Markdown formats
        import re
        
        # Handle bold **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    
    def _generate_html(
        self, 
        content: str, 
        output_path: str, 
        style_config: Dict
    ) -> Dict[str, Any]:
        """Generate HTML document."""
        colors = style_config.get("colors", self.DEFAULT_STYLE["colors"])
        
        # Simple HTML template
        html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Research Report</title>
    <style>
        body {{
            font-family: Georgia, serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{ color: #{colors['primaryDark']}; border-bottom: 2px solid #{colors['accent']}; }}
        h2 {{ color: #{colors['primary']}; margin-top: 30px; }}
        h3 {{ color: #{colors['primary']}; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
        blockquote {{ border-left: 4px solid #{colors['accent']}; padding-left: 15px; color: #666; }}
    </style>
</head>
<body>
{self._markdown_to_html(content)}
</body>
</html>"""
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # Save HTML
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        file_size = os.path.getsize(output_path)
        
        return {
            "success": True,
            "file_path": output_path,
            "file_size": file_size,
            "format": "html",
            "pages": "N/A",
        }
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert Markdown to HTML.
        
        Simplified implementation.
        """
        import re
        
        html = content
        
        # Headings
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        
        # Bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        
        # Italic
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Horizontal rule
        html = re.sub(r'^---$', r'<hr>', html, flags=re.MULTILINE)
        
        # Paragraphs
        paragraphs = html.split('\n\n')
        html_paragraphs = []
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith('<'):
                p = f'<p>{p}</p>'
            html_paragraphs.append(p)
        
        return '\n\n'.join(html_paragraphs)
    
    def _save_markdown(self, content: str, output_path: str) -> Dict[str, Any]:
        """Save Markdown file."""
        # Ensure output directory exists
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        file_size = os.path.getsize(output_path)
        
        return {
            "success": True,
            "file_path": output_path,
            "file_size": file_size,
            "format": "markdown",
            "pages": "N/A",
        }
