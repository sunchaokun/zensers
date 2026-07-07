# -*- coding: utf-8 -*-
"""
Document Generation Agent
=========================

Unified document generation agent supporting multiple output formats and version control.

Design Principles:
1. Inherits from FixedAgent base class
2. Supports multiple document formats (Word/PPT/PDF)
3. Supports multiple actions (generate/rollback/compare/export)
4. Session context integration
5. Comprehensive error handling
6. LLM content revision support (Week 34)

Architecture Reference: docs/UNIFIED_DOCUMENT_GENERATION_AGENT_DESIGN.md
"""

import json
import logging
import re
import shutil
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from .base_fixed_agent import FixedAgent
from .document_models import (
    DocumentFormat,
    GenerationAction,
    DocumentGenerationRequest,
    DocumentGenerationResult,
    DocumentVersion,
    ValidationError,
)

from src.core.adjustment import SectionLocator, ContentApplier
from src.core.llm_client import call_llm

logger = logging.getLogger(__name__)

# ==================== Constants ====================
MAX_REVISION_CONTENT_LENGTH = 5000  # Max length for revision content storage
MAX_PROMPT_CONTENT_LENGTH = 3000    # Max original content length in prompt
DEFAULT_LLM_MAX_TOKENS = 4096       # Default LLM max tokens
DEFAULT_LLM_TEMPERATURE = 0.3       # Default LLM temperature (lower for consistency in revisions)


class DocumentGenerationAgent(FixedAgent):
    """
    Document Generation Agent
    
    Responsible for converting research results into professional format document outputs.
    Supports Word/PPT/PDF formats with version control capabilities.
    
    Core Capabilities:
    - Multi-format document generation
    - Version management (create/list/compare/rollback)
    - Export management
    - Preview generation
    - Content adjustment
    
    Attributes:
        agent_type: Agent type identifier
        version: Agent version number
        capabilities: Agent capability list
    """
    
    agent_type = "document_generation"
    version = "1.0.0"
    capabilities = [
        "Word document generation",
        "PPT document generation",
        "PDF document generation",
        "Version management",
        "Export management",
        "Preview generation",
        "Content adjustment"
    ]
    
    def __init__(
        self,
        agent_id: str,
        name: str = "DocumentGenerationAgent",
        description: str = "Unified Document Generation Agent",
        storage_path: Optional[str] = None,
    ):
        """
        Initialize Document Generation Agent
        
        Args:
            agent_id: Agent unique identifier
            name: Agent name
            description: Agent description
            storage_path: Storage path
        """
        super().__init__(agent_id, name=name, description=description, storage_path=storage_path)
        
        # Storage path (with exception handling)
        self._storage_path = Path(storage_path) if storage_path else Path("data")
        self._documents_dir = self._storage_path / "documents"
        
        try:
            self._documents_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            raise RuntimeError(f"Failed to create documents directory: {e}") from e
        
        # Session integration
        self._shared_memory: Optional[Any] = None
        self._message_bus: Optional[Any] = None
        # RG-FIX-2: cross-section deduplication fingerprint set
        self._global_seen_paragraphs: set = set()
        
        # Revision manager
        self._revision_manager: Optional[Any] = None
    
    def set_shared_memory(self, shared_memory: Any) -> None:
        """
        Set SharedMemory instance
        
        Args:
            shared_memory: SharedMemory instance
        """
        self._shared_memory = shared_memory
    
    def set_message_bus(self, message_bus: Any) -> None:
        """
        Set MessageBus instance
        
        Args:
            message_bus: MessageBus instance
        """
        self._message_bus = message_bus
    
    def set_revision_manager(self, revision_manager: Any) -> None:
        """
        Set revision manager
        
        Args:
            revision_manager: RevisionManager instance
        """
        self._revision_manager = revision_manager
    
    def _validate_document_path(self, document_path: str) -> Path:
        """
        Validate document path security
        
        Prevents path traversal attacks, ensuring path is within allowed directories.
        
        Args:
            document_path: Document path to validate
            
        Returns:
            Validated Path object
            
        Raises:
            ValueError: Path is not in allowed directory or path is empty
        """
        if not document_path:
            raise ValueError("Document path cannot be empty")
        
        path = Path(document_path).resolve()
        allowed_dir = self._documents_dir.resolve()
        
        # Check for dangerous path patterns
        dangerous_patterns = ['../', '..\\', '/etc/', '/root/', '\\Windows\\']
        for pattern in dangerous_patterns:
            if pattern in document_path:
                logger.warning(f"Dangerous path pattern detected: {document_path}")
                raise ValueError(f"Document path contains dangerous pattern: {pattern}")
        
        # Allowed directory check (relaxed mode, allows output and temp directories)
        allowed_prefixes = [
            str(allowed_dir),
            str(Path("output").resolve()),
            str(Path("data").resolve()),
            str(Path("temp").resolve()),
        ]
        
        path_str = str(path)
        for prefix in allowed_prefixes:
            if path_str.startswith(prefix):
                return path
        
        # If not in allowed directory, log warning but allow (backward compatibility)
        logger.warning(f"Document path outside recommended directories: {document_path}")
        return path
    
    def validate_input(self, task_input: Dict[str, Any]) -> tuple[bool, str]:
        """
        Validate input parameters
        
        Args:
            task_input: Input parameters to validate
            
        Returns:
            (is_valid, error_message)
        """
        # Basic validation
        valid, error = super().validate_input(task_input)
        if not valid:
            return valid, error
        
        # Validate action field
        if "action" not in task_input:
            return False, "Missing required field: action"
        
        # Validate action value
        action_value = task_input["action"]
        valid_actions = [a.value for a in GenerationAction]
        if action_value not in valid_actions:
            return False, f"Invalid action: {action_value}. Valid actions: {valid_actions}"
        
        # Validate required fields based on action
        try:
            action = GenerationAction(action_value)
        except ValueError:
            return False, f"Invalid action: {action_value}"
        
        # produce_document and regenerate_document require output_format
        if action in [GenerationAction.PRODUCE_DOCUMENT, GenerationAction.REGENERATE_DOCUMENT]:
            if "output_format" not in task_input:
                return False, "Missing required field: output_format for produce/regenerate action"
            
            # Validate output_format value
            valid_formats = [f.value for f in DocumentFormat]
            if task_input["output_format"] not in valid_formats:
                return False, f"Invalid output_format: {task_input['output_format']}. Valid formats: {valid_formats}"
        
        # rollback_version requires task_id and version_id
        if action == GenerationAction.ROLLBACK_VERSION:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for rollback action"
            if "version_id" not in task_input:
                return False, "Missing required field: version_id for rollback action"
        
        # compare_versions requires task_id and two version_ids
        if action == GenerationAction.COMPARE_VERSIONS:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for compare action"
            if "version_id" not in task_input or "version_id_2" not in task_input:
                return False, "Missing required fields: version_id and version_id_2 for compare action"
        
        # list_versions requires task_id
        if action == GenerationAction.LIST_VERSIONS:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for list_versions action"
        
        # export_document requires task_id and export_path
        if action == GenerationAction.EXPORT_DOCUMENT:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for export action"
            if "export_path" not in task_input:
                return False, "Missing required field: export_path for export action"
        
        # get_preview requires task_id
        if action == GenerationAction.GET_PREVIEW:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for preview action"
        
        # adjust_content requires task_id and adjustments
        if action == GenerationAction.ADJUST_CONTENT:
            if "task_id" not in task_input:
                return False, "Missing required field: task_id for adjust action"
        
        return True, ""
    
    async def execute(self, task_input: Dict[str, Any]) -> Dict[str, Any]:
        """
        Execute document generation task (asynchronous)

        Dispatches to corresponding handler based on action type.

        Args:
            task_input: Task input parameters

        Returns:
            Task execution result
        """
        # Publish start event
        await self.publish_event("document_generation_started", task_input.get("action", "unknown"))
        
        # Parse request
        try:
            request = DocumentGenerationRequest.from_dict(task_input)
        except ValidationError as e:
            return DocumentGenerationResult(
                success=False,
                error=str(e),
                error_code="VALIDATION_ERROR"
            ).to_dict()
        
        # Dispatch based on action
        action_handlers = {
            GenerationAction.PRODUCE_DOCUMENT: self._handle_produce_document,
            GenerationAction.REGENERATE_DOCUMENT: self._handle_regenerate_document,
            GenerationAction.ADJUST_CONTENT: self._handle_adjust_content,
            GenerationAction.EXPORT_DOCUMENT: self._handle_export_document,
            GenerationAction.GET_PREVIEW: self._handle_get_preview,
            GenerationAction.LIST_VERSIONS: self._handle_list_versions,
            GenerationAction.ROLLBACK_VERSION: self._handle_rollback_version,
            GenerationAction.COMPARE_VERSIONS: self._handle_compare_versions,
        }
        
        handler = action_handlers.get(request.action)
        if not handler:
            return DocumentGenerationResult(
                success=False,
                error=f"Unsupported action: {request.action}",
                error_code="UNSUPPORTED_ACTION"
            ).to_dict()
        
        import asyncio as _asyncio
        if _asyncio.iscoroutinefunction(handler):
            result = await handler(request)
        else:
            result = handler(request)
        
        # Publish completion event
        await self.publish_event("document_generation_completed", {"success": result.get("success", False)})
        
        return result
    
    def _handle_produce_document(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle document generation request

        Uses JSON → HTML → Word/PPT workflow to generate document:
        1. Get content from research result or history
        2. Use ContentOrchestrator to convert to HTML
        3. Use HTMLToWordConverter/HTMLToPPTConverter to generate final document

        Args:
            request: Document generation request

        Returns:
            Generation result
        """
        # Generate task ID
        task_id = request.task_id or f"doc_{uuid.uuid4().hex[:8]}"
        
        # 1. Get research content
        research_result = request.research_result
        if not research_result:
            # Try to load from history
            research_result = self._load_research_result(task_id)
        
        if not research_result:
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error="No research result provided and no historical record found",
                error_code="NO_CONTENT"
            ).to_dict()
        
        # 2. Determine output format
        output_format = request.output_format or DocumentFormat.DOCX
        format_value = output_format.value if hasattr(output_format, 'value') else str(output_format)
        
        # 3. Generate output path (prefer directory specified in request)
        if request.output_dir:
            output_dir = Path(request.output_dir)
        else:
            output_dir = self._documents_dir / "generated"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{task_id}_report_{timestamp}.{format_value}"
        output_path = output_dir / output_filename
        
        # 4. Use document generation workflow
        try:
            # HTML preview: use ContentOrchestrator (not DocumentGenerator which only does DOCX/PDF)
            if format_value == "html":
                # 前置图表生成：从 data_points 和段落内容自动配图
                research_result = self._generate_charts_for_html(research_result)

                from src.content.content_orchestrator import ContentOrchestrator
                from src.core.preview_storage import PreviewStorage
                orchestrator = ContentOrchestrator()
                html_content = orchestrator.transform_to_html(
                    research_result=research_result,
                    output_format="html",
                    output_dir=str(PreviewStorage.NEW_DIR),
                )
                output_path.write_text(html_content, encoding="utf-8")
                
                file_size = output_path.stat().st_size
                pages_estimate = max(1, len(research_result.get("sections", [])))
                logger.info(f"HTML preview generated: {output_path}")
                
                return DocumentGenerationResult(
                    success=True,
                    task_id=task_id,
                    output_format=output_format,
                    version_id="v1",
                    document_path=str(output_path),
                    file_size=file_size,
                    pages_estimate=pages_estimate,
                    message=f"Preview generated: {output_filename}"
                ).to_dict()

            # DOCX/PDF/PPTX: use DocumentGenerator
            from ...core.orchestrator.output.document_generator import (
                DocumentGenerator,
                DocumentConfig,
                DocumentFormat as DGFormat,
            )
            
            # Map format
            format_map = {
                DocumentFormat.DOCX: DGFormat.DOCX,
                DocumentFormat.PDF: DGFormat.PDF,
                DocumentFormat.PPTX: DGFormat.PPTX,
            }
            dg_format = format_map.get(output_format, DGFormat.DOCX)
            
            # Create document config
            config = DocumentConfig(
                format=dg_format,
                title=research_result.get("title") or research_result.get("topic") or "Research Report",
                author=research_result.get("author", "AI Research Assistant"),
                template_path=Path(request.template) if request.template else None,
            )
            
            # Create document generator
            generator = DocumentGenerator(config)
            
            # Add content
            self._populate_document_content(generator, research_result)
            
            # Generate document
            result = generator.generate(output_path)
            
            if result.path and result.path.exists():
                file_size = result.path.stat().st_size
                pages_estimate = max(1, len(research_result.get("sections", [])))
                
                logger.info(f"Document generated successfully: {output_path}")
                
                return DocumentGenerationResult(
                    success=True,
                    task_id=task_id,
                    output_format=output_format,
                    version_id="v1",
                    document_path=str(output_path),
                    file_size=file_size,
                    pages_estimate=pages_estimate,
                    message=f"Document generated: {output_filename}"
                ).to_dict()
            else:
                return DocumentGenerationResult(
                    success=False,
                    task_id=task_id,
                    error="Document generation failed - no output file created",
                    error_code="GENERATION_FAILED"
                ).to_dict()
                
        except ImportError as e:
            logger.warning(f"DocumentGenerator not available: {e}, using fallback")
            return self._fallback_generate_document(
                task_id=task_id,
                research_result=research_result,
                output_format=format_value,
                output_path=output_path,
            )
        except Exception as e:
            logger.error(f"Document generation error: {e}", exc_info=True)
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error=f"Document generation failed: {str(e)}",
                error_code="GENERATION_ERROR"
            ).to_dict()
    
    def _load_research_result(self, task_id: str) -> Optional[Dict[str, Any]]:
        """
        Load research result from history

        Args:
            task_id: Task ID

        Returns:
            Research result dictionary, or None if not found
        """
        # Try to load from SharedMemory
        if self._shared_memory:
            try:
                result = self._shared_memory.get(f"research_result_{task_id}")
                if result:
                    return result
            except Exception as e:
                logger.warning(f"Failed to load from SharedMemory: {e}")
        
        # Try to load from file
        result_file = self._documents_dir.parent / "research_results" / f"{task_id}.json"
        if result_file.exists():
            try:
                import json
                with open(result_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Extract research result
                    if "result" in data and "data" in data["result"]:
                        return data["result"]["data"]
                    return data
            except Exception as e:
                logger.warning(f"Failed to load research result from file: {e}")
        
        return None
    
    def _populate_document_content(
        self,
        generator: Any,
        research_result: Dict[str, Any]
    ) -> None:
        """
        Populate research result into document generator

        Args:
            generator: DocumentGenerator instance
            research_result: Research result dictionary
        """
        # [FIX Bug 2] Store original section metadata (data_points, charts) for chart generation
        sections_data = research_result.get("sections", [])
        if hasattr(generator, 'set_sections_meta'):
            generator.set_sections_meta(sections_data)

        # Add title
        title=research_result.get("title", "研究报告")
        generator.add_heading(title, level=0)
        
        # Add metadata
        if research_result.get("author"):
            generator.add_paragraph(f"Author: {research_result['author']}")
        if research_result.get("date"):
            generator.add_paragraph(f"Date: {research_result['date']}")
        
        generator.add_paragraph("")  # Empty line

        # Add sections
        sections = research_result.get("sections", [])
        for section in sections:
            section_title = section.get("title", "")
            section_content = section.get("content", "")
            section_data_points = section.get("data_points", [])
            
            if section_title:
                generator.add_heading(section_title, level=1)
            
            if section_content:
                # Clean content: remove LLM conversation prefixes and irrelevant content
                cleaned_content = self._clean_llm_content(section_content)
                
                # Remove duplicate title at content start (avoid "Executive Summary" appearing twice)
                cleaned_content = self._strip_duplicate_title(cleaned_content, section_title)
                
                # **New**: Remove duplicate paragraphs (LLM may generate duplicate content)
                cleaned_content = self._deduplicate_paragraphs(cleaned_content, global_dedup=True)  # RG-FIX-2
                
                # Process Markdown format content
                elements = self._process_markdown_content(cleaned_content)
                
                for element in elements:
                    if element["type"] == "heading":
                        # Skip headings in content: section.title already added as heading
                        continue
                    elif element["type"] == "paragraph":
                        generator.add_paragraph(element["text"])
                    elif element["type"] == "list":
                        for item in element["items"]:
                            generator.add_paragraph(f"• {item}")
                    elif element["type"] == "ordered_list":
                        for j, item in enumerate(element["items"], 1):
                            generator.add_paragraph(f"{j}. {item}")
                    elif element["type"] == "table":
                        if len(element["rows"]) > 1:
                            headers = element["rows"][0]
                            rows = element["rows"][1:]
                                # Filter out separator rows (like |---|---|) and empty value rows
                            rows = [r for r in rows if not all(
                                (c is None) or (c.replace("-", "").replace(":", "") == "") 
                                for c in r
                            )]
                            if rows:
                                generator.add_table(headers, rows)
                
                # Smart chart generation: automatically add charts for suitable sections
                if self._should_generate_charts(section_title):
                    try:
                        # [FIX Bug 3] Use original section_content (not cleaned) for data extraction
                        # so market share, growth rates, etc. aren't lost during dedup/stripping
                        chart_paths = generator.add_smart_chart(
                            section_title=section_title,
                            content=section_content,
                            data_points=section_data_points,
                            max_charts=2
                        )
                        if chart_paths:
                            logger.info(f"Auto-generated {len(chart_paths)} charts for section: {section_title}")
                            # [FIX Bug 1] Embed generated chart images into document content
                            for chart_path in chart_paths:
                                try:
                                    generator.add_image(Path(chart_path))
                                except Exception as img_e:
                                    logger.warning(f"Failed to add chart image {chart_path}: {img_e}")
                    except Exception:
                        logger.exception("Failed to generate smart charts")
            
            # P0-3 fix: Handle charts in section
            section_charts = section.get("charts", [])
            for chart in section_charts:
                chart_path = chart.get("path", "")
                if chart_path and Path(chart_path).exists():
                    try:
                        generator.add_image(chart_path)
                        logger.info(f"Added chart image: {chart_path}")
                    except Exception as e:
                        logger.warning(f"Failed to add chart image {chart_path}: {e}")
            
            # Process subsections
            subsections = section.get("subsections", [])
            for subsection in subsections:
                sub_title = subsection.get("title", "")
                sub_content = subsection.get("content", "")
                sub_points = subsection.get("points", []) or []
                
                if sub_title:
                    generator.add_heading(sub_title, level=2)
                
                if sub_points:
                    for pt in sub_points:
                        generator.add_heading(pt, level=3)
                        pt_content = self._extract_point_text(sub_content, pt)
                        if pt_content:
                            cleaned = self._clean_llm_content(pt_content)
                            elements = self._process_markdown_content(cleaned)
                            for element in elements:
                                if element["type"] == "heading":
                                    continue
                                elif element["type"] == "paragraph":
                                    generator.add_paragraph(element["text"])
                                elif element["type"] == "list":
                                    for item in element["items"]:
                                        generator.add_paragraph(f"• {item}")
                                elif element["type"] == "ordered_list":
                                    for j, item in enumerate(element["items"], 1):
                                        generator.add_paragraph(f"{j}. {item}")
                                elif element["type"] == "table":
                                    if len(element["rows"]) > 1:
                                        headers = element["rows"][0]
                                        rows = element["rows"][1:]
                                        rows = [r for r in rows if not all(
                                            (c is None) or (c.replace("-", "").replace(":", "") == "")
                                            for c in r
                                        )]
                                        if rows:
                                            generator.add_table(headers, rows)
                elif sub_content:
                    cleaned_content = self._clean_llm_content(sub_content)
                    elements = self._process_markdown_content(cleaned_content)
                    
                    for element in elements:
                        if element["type"] == "heading":
                            heading_text = element.get("text", "").strip()
                            if heading_text == sub_title:
                                continue
                            generator.add_heading(heading_text, level=3)
                        elif element["type"] == "paragraph":
                            generator.add_paragraph(element["text"])
                        elif element["type"] == "list":
                            for item in element["items"]:
                                generator.add_paragraph(f"• {item}")
                        elif element["type"] == "ordered_list":
                            for j, item in enumerate(element["items"], 1):
                                generator.add_paragraph(f"{j}. {item}")
                        elif element["type"] == "table":
                            if len(element["rows"]) > 1:
                                headers = element["rows"][0]
                                rows = element["rows"][1:]
                                rows = [r for r in rows if not all(
                                    (c is None) or (c.replace("-", "").replace(":", "") == "") 
                                    for c in r
                                )]
                                if rows:
                                    generator.add_table(headers, rows)
        
        # Add key findings
        key_findings = research_result.get("key_findings", [])
        if key_findings:
            generator.add_heading("Key Findings", level=1)
            for finding in key_findings:
                generator.add_paragraph(f"• {finding}")
        
        # Add data points
        data_points = research_result.get("data_points", [])
        if data_points:
            generator.add_heading("Key Data", level=1)
            headers = ["Metric", "Value", "Unit"]
            rows = [
                [dp.get("metric", ""), dp.get("value", ""), dp.get("unit", "")]
                for dp in data_points
            ]
            generator.add_table(headers, rows)
        
        # P0-3 fix: Add data sources/references
        sources = research_result.get("sources", [])
        if sources:
            generator.add_heading("Data Sources", level=1)
            for i, source in enumerate(sources[:20], 1):  # Show at most 20 sources
                title = source.get("title", "")
                url = source.get("url", "")
                if title and url:
                    generator.add_paragraph(f"[{i}] {title} - {url}")
                elif url:
                    generator.add_paragraph(f"[{i}] {url}")
                elif title:
                    generator.add_paragraph(f"[{i}] {title}")
    
    @staticmethod
    def _extract_point_text(full_content: str, point_title: str) -> str:
        """Extract text belonging to a specific point from subsection content.
        
        Searches for the point heading in markdown content and returns the
        text between this heading and the next heading or end of content.
        """
        if not full_content or not point_title:
            return ""
        import re
        escaped = re.escape(point_title)
        pattern = rf'^#{{1,4}}\s+{escaped}\s*$'
        lines = full_content.split('\n')
        capturing = False
        captured = []
        for line in lines:
            if re.match(pattern, line.strip(), re.IGNORECASE):
                capturing = True
                continue
            if capturing:
                if re.match(r'^#{1,4}\s+', line.strip()):
                    break
                captured.append(line)
        return '\n'.join(captured).strip()
    
    def _strip_duplicate_title(self, content: str, section_title: str) -> str:
        """
        Remove duplicate title at content start to avoid title being rendered twice

        Scenario: section title is already "Executive Summary",
        and content starts with "## Executive Summary", "### Executive Summary", or "**Executive Summary**",
        causing the title to appear twice.

        Args:
            content: Original content
            section_title: Section title

        Returns:
            Cleaned content
        """
        if not content or not section_title:
            return content
        
        import re
        
        lines = content.split("\n")
        cleaned = []
        title_stripped = section_title.strip()
        
        # Build patterns to skip: various markdown forms of the title
        skip_patterns = [
            r'^#{1,6}\s*' + re.escape(title_stripped) + r'\s*$',
            r'^\*\*' + re.escape(title_stripped) + r'\*\*\s*$',
            r'^' + re.escape(title_stripped) + r'\s*$',
        ]
        
        for i, line in enumerate(lines):
            stripped = line.strip()
                # Extend scope to first 20 lines, covering more LLM output prefixes
            if i < 20:
                should_skip = False
                for pattern in skip_patterns:
                    if re.match(pattern, stripped):
                        should_skip = True
                        break
                if should_skip:
                    continue
            cleaned.append(line)
        
        return "\n".join(cleaned)
    
    def _deduplicate_paragraphs(self, content: str, global_dedup: bool = False) -> str:
        """
        Remove duplicate paragraphs from content (RG-FIX-2: cross-section support)

        LLM sometimes generates duplicate paragraphs (e.g., same sentence output twice).
        This method detects and removes consecutive duplicate paragraphs.
        When global_dedup=True, uses class-level _global_seen_paragraphs for cross-section dedup.

        Args:
            content: Original content
            global_dedup: If True, check against global fingerprint set

        Returns:
            Deduplicated content
        """
        if not content:
            return content
        
        content = self._clean_duplicate_html_tags(content)
        content = self._clean_inline_duplicates(content)
        
        lines = content.split('\n')
        result = []
        prev_paragraph = []
        prev_paragraph_text = ""
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            if stripped:
                prev_paragraph.append(stripped)
                i += 1
                continue
            
            if prev_paragraph:
                current_text = ' '.join(prev_paragraph)
                
                # RG-FIX-2: cross-section dedup via global fingerprint
                if global_dedup:
                    _fingerprint = current_text.strip()[:80]
                    if _fingerprint in self._global_seen_paragraphs:
                        logger.debug(f"RG-FIX-2: Skipping cross-section duplicate: {current_text[:50]}...")
                        prev_paragraph = []
                        i += 1
                        continue
                    self._global_seen_paragraphs.add(_fingerprint)
                
                if self._is_similar_paragraph(current_text, prev_paragraph_text):
                    logger.debug(f"Skipping duplicate paragraph: {current_text[:50]}...")
                else:
                    result.extend(prev_paragraph)
                    result.append('')
                    prev_paragraph_text = current_text
                
                prev_paragraph = []
            
            i += 1
        
        if prev_paragraph:
            current_text = ' '.join(prev_paragraph)
            if global_dedup:
                _fingerprint = current_text.strip()[:80]
                if _fingerprint in self._global_seen_paragraphs:
                    return '\n'.join(result)
                self._global_seen_paragraphs.add(_fingerprint)
            if not self._is_similar_paragraph(current_text, prev_paragraph_text):
                result.extend(prev_paragraph)
        
        return '\n'.join(result)
    
    def _clean_duplicate_html_tags(self, content: str) -> str:
        """
        Clean duplicate HTML tags

        Handles duplicate tags that LLM might generate, such as:
        - <table><table> → <table>
        - <thead><tr><thead><tr> → <thead><tr>
        - </tbody></table></tbody></table> → </tbody></table>

        Args:
            content: Original content

        Returns:
            Cleaned content
        """
        import re
        
        # Common duplicate tag patterns
        duplicate_patterns = [
            # Opening tag duplicates
            (r'(<table[^>]*>)\s*\1', r'\1'),
            (r'(<thead[^>]*>)\s*\1', r'\1'),
            (r'(<tbody[^>]*>)\s*\1', r'\1'),
            (r'(<tr[^>]*>)\s*\1', r'\1'),
            (r'(<td[^>]*>)\s*\1', r'\1'),
            (r'(<th[^>]*>)\s*\1', r'\1'),
            (r'(<ul[^>]*>)\s*\1', r'\1'),
            (r'(<ol[^>]*>)\s*\1', r'\1'),
            (r'(<li[^>]*>)\s*\1', r'\1'),
            (r'(<div[^>]*>)\s*\1', r'\1'),
            (r'(<section[^>]*>)\s*\1', r'\1'),
            # Closing tag duplicates
            (r'(</table>)\s*\1', r'\1'),
            (r'(</thead>)\s*\1', r'\1'),
            (r'(</tbody>)\s*\1', r'\1'),
            (r'(</tr>)\s*\1', r'\1'),
            (r'(</td>)\s*\1', r'\1'),
            (r'(</th>)\s*\1', r'\1'),
            (r'(</ul>)\s*\1', r'\1'),
            (r'(</ol>)\s*\1', r'\1'),
            (r'(</li>)\s*\1', r'\1'),
            (r'(</div>)\s*\1', r'\1'),
            (r'(</section>)\s*\1', r'\1'),
        ]
        
        for pattern, replacement in duplicate_patterns:
            # Multiple replacements until no change
            prev_content = None
            while prev_content != content:
                prev_content = content
                content = re.sub(pattern, replacement, content, flags=re.IGNORECASE)
        
        return content
    
    def _clean_inline_duplicates(self, content: str) -> str:
        """
        Clean inline duplicate text

        Handles duplicate text that LLM might generate, such as:
        - "市场概况市场概况" → "市场概况"
        - "研究结论研究结论" → "研究结论"

        Args:
            content: Original content

        Returns:
            Cleaned content
        """
        import re
        
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                cleaned_lines.append(line)
                continue
            
            # Detect and fix inline duplicates (word/phrase level duplicates)
            # Pattern: detect consecutive duplicate text fragments (length >= 4 characters)
            # Example: "市场概况市场概况" → "市场概况"

            # Special handling for Chinese text
            if len(stripped) >= 8:  # Only process lines long enough
                # Check if first half and second half are identical
                half_len = len(stripped) // 2
                if half_len >= 4:
                    first_half = stripped[:half_len]
                    second_half = stripped[half_len:]
                    
                    # If both halves are identical or highly similar
                    if first_half == second_half:
                        stripped = first_half
                        logger.debug(f"Cleaned inline duplicate: {line.strip()[:50]}...")
                    elif self._is_similar_paragraph(first_half, second_half, threshold=0.95):
                        stripped = first_half
                        logger.debug(f"Cleaned similar inline duplicate: {line.strip()[:50]}...")
            
            # Detect duplicates in Markdown heading lines
            # Example: "## 市场概况## 市场概况" → "## 市场概况"
            title_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if title_match:
                prefix = title_match.group(1)
                title_text = title_match.group(2)
                
                # Check if title text is duplicated
                half_len = len(title_text) // 2
                if half_len >= 2:
                    first_half = title_text[:half_len]
                    second_half = title_text[half_len:]
                    if first_half == second_half:
                        stripped = f"{prefix} {first_half}"
                        logger.debug(f"Cleaned title duplicate: {line.strip()[:50]}...")
            
            # Detect duplicates in table rows
            # Example: "| 指标 | 数值 || 指标 | 数值 |" → "| 指标 | 数值 |"
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.split("|")[1:-1] if c.strip()]
                if cells:
                    half_len = len(cells) // 2
                    if half_len >= 1:
                        first_half = cells[:half_len]
                        second_half = cells[half_len:]
                        if first_half == second_half:
                            stripped = "| " + " | ".join(first_half) + " |"
                            logger.debug(f"Cleaned table row duplicate: {line.strip()[:50]}...")
            
            cleaned_lines.append(stripped)
        
        return '\n'.join(cleaned_lines)
    
    def _is_similar_paragraph(self, text1: str, text2: str, threshold: float = 0.9) -> bool:
        """
        Determine if two paragraphs are similar

        Args:
            text1: First paragraph
            text2: Second paragraph
            threshold: Similarity threshold (0-1)

        Returns:
            Whether similar
        """
        if not text1 or not text2:
            return False
        
        # Exactly identical
        if text1 == text2:
            return True
        
        # Compare after normalization
        import re
        norm1 = re.sub(r'\s+', '', text1)
        norm2 = re.sub(r'\s+', '', text2)
        
        if norm1 == norm2:
            return True
        
# Calculate similarity (simple character overlap rate)
        if len(norm1) == 0 or len(norm2) == 0:
            return False
        
        # Use longer text as baseline
        shorter, longer = (norm1, norm2) if len(norm1) < len(norm2) else (norm2, norm1)
        
        # Check if shorter text is a substring of longer text
        if shorter in longer:
            overlap_ratio = len(shorter) / len(longer)
            return overlap_ratio >= threshold
        
        return False
    
    def _clean_llm_content(self, content: str) -> str:
        """
        Clean LLM-generated content, remove conversation prefixes and irrelevant content

        Args:
            content: Original content

        Returns:
            Cleaned content
        """
        if not content:
            return content
        
        # Define conversation prefix patterns to remove
        prefixes_to_remove = [
            "好的，收到您的研究任务",
            "好的，收到您的综合分析任务",
            "好的，收到您的深度研究任务",
            "好的，收到您的",
            "好的，基于您提供的",
            "好的，我将",
            "好的，让我",
            "好的，这是",
            "收到，我将",
            "收到，让我",
        ]
        
        # Define prompt trace patterns to delete entire lines
        # **Fix**: Add more cleaning patterns
        prompt_patterns_to_remove = [
            r'^基于您提供的.*数据点.*分析.*',
            r'^基于您提供的所有数据点.*',
            r'.*将对["""].*["""].*维度进行.*',
            r'.*形成一份面向决策层.*',
            r'.*原创性的综合与提炼.*',
            r'^数据来源\s*数据来源\s*$',  # Duplicate data source header
            r'^原创洞察[：:]\s*',        # Remove "Original insight:" at line start
            r'^原创洞察\s*',              # Remove "Original insight" at line start
            r'.*基于多源数据.*',          # Remove data source description
            r'^本.*基于.*数据.*分析.*',   # Remove "This executive summary is based on multi-source data..."
            r'原创洞察[：:]\s*',          # Remove "Original insight:" in paragraph
            r'原创洞察\s*[：:]',          # Remove "Original insight:" (with space variant)
        ]
        
        # Check and remove prefixes
        lines = content.split("\n")
        cleaned_lines = []
        skip_initial_prefix = True
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # Skip initial conversation prefixes
            if skip_initial_prefix and i < 5:  # 只检查前5行
                is_prefix = False
                for prefix in prefixes_to_remove:
                    if stripped.startswith(prefix):
                        is_prefix = True
                        break
                if is_prefix:
                    continue
                else:
                    skip_initial_prefix = False
            
            # Skip various separator lines (usually LLM reply format markers)
            # **Fix**: Extend separator cleaning, cover more forms, remove line count limit
            if re.match(r'^[-*_]{3,}$', stripped):
                continue
            
            # Check if matches prompt trace pattern to delete
            should_remove = False
            for pattern in prompt_patterns_to_remove:
                if re.match(pattern, stripped):
                    should_remove = True
                    break
            if should_remove:
                continue
            
            # **Fix**: Replace inline patterns (like "Original insight:")
            cleaned_line = line
            for pattern in [r'原创洞察[：:]\s*', r'原创洞察\s*[：:]']:
                cleaned_line = re.sub(pattern, '', cleaned_line)
            
            cleaned_lines.append(cleaned_line)
        
        return "\n".join(cleaned_lines)
    
    def _process_markdown_content(self, content: str) -> List[Dict[str, Any]]:
        """
        Process Markdown format content, extract structured elements

        Supported formats:
        - Headings: # ~ ######
        - Lists: - or *
        - Ordered lists: 1. 2. 3.
        - Tables: | col1 | col2 |
        - Bold: **text**
        - Italic: *text* or _text_
        - Code: `code`

        Args:
            content: Markdown format content

        Returns:
            List of structured elements
        """
        import re
        elements = []
        lines = content.split("\n")
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Heading processing
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                title = self._clean_inline_markdown(heading_match.group(2))
                elements.append({
                    "type": "heading",
                    "level": level,
                    "text": title
                })
                i += 1
                continue
            
            # List item processing
            list_match = re.match(r'^[-*]\s+(.+)$', stripped)
            if list_match:
                items = [self._clean_inline_markdown(list_match.group(1))]
                i += 1
                # Collect consecutive list items
                while i < len(lines):
                    next_match = re.match(r'^[-*]\s+(.+)$', lines[i].strip())
                    if next_match:
                        items.append(self._clean_inline_markdown(next_match.group(1)))
                        i += 1
                    else:
                        break
                elements.append({
                    "type": "list",
                    "items": items
                })
                continue
            
            # Ordered list processing
            ordered_match = re.match(r'^\d+\.\s+(.+)$', stripped)
            if ordered_match:
                items = [self._clean_inline_markdown(ordered_match.group(1))]
                i += 1
                # Collect consecutive ordered list items
                while i < len(lines):
                    next_match = re.match(r'^\d+\.\s+(.+)$', lines[i].strip())
                    if next_match:
                        items.append(self._clean_inline_markdown(next_match.group(1)))
                        i += 1
                    else:
                        break
                elements.append({
                    "type": "ordered_list",
                    "items": items
                })
                continue
            
            # Table processing (simple tables)
            if stripped.startswith("|") and "|" in stripped[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) > 2:  # 至少有表头和一行数据
                    # Parse table
                    rows = []
                    for tl in table_lines:
                        cells = [self._clean_inline_markdown(c.strip()) for c in tl.split("|")[1:-1]]
                        rows.append(cells)
                    elements.append({
                        "type": "table",
                        "rows": rows
                    })
                continue
            
            # Normal paragraph
            if stripped:
                paragraph_lines = [stripped]
                i += 1
                # Collect consecutive non-empty lines as paragraph
                while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6}|\*|[-*]|\d+\.)', lines[i].strip()):
                    paragraph_lines.append(lines[i].strip())
                    i += 1
                # Clean inline Markdown formatting
                cleaned_text = self._clean_inline_markdown(" ".join(paragraph_lines))
                elements.append({
                    "type": "paragraph",
                    "text": cleaned_text
                })
            else:
                i += 1
        
        return elements
    
    def _clean_inline_markdown(self, text: str) -> str:
        """
        Clean inline Markdown formatting, convert to plain text

        Handles:
        - **bold** → bold
        - *italic* → italic
        - `code` → code
        - [link](url) → link
        - ~~strikethrough~~ → strikethrough

        Args:
            text: Text containing Markdown formatting

        Returns:
            Cleaned plain text
        """
        import re
        
        # Handle bold **text** or __text__
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        
        # Handle italic *text* or _text_
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        
        # Handle code `code`
        text = re.sub(r'`(.+?)`', r'\1', text)
        
        # Handle link [text](url)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        
        # Handle strikethrough ~~text~~
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        
        # Handle HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        
        return text

    # ==================== 图表系统修复：HTML 路径自动配图 ====================

    def _generate_charts_for_html(self, research_result: Dict[str, Any]) -> Dict[str, Any]:
        """为 HTML 报告生成图表：三通道策略"""
        if not research_result.get("sections"):
            return research_result

        self._html_charts_from_datapoints(research_result)
        self._html_charts_from_content(research_result)

        return research_result

    # 指标-数值语义校验白名单
    _METRIC_SEMANTIC_RULES = {
        "净利润": {"max_value": 10000, "unit_hint": "亿元"},
        "营收": {"max_value": 100000, "unit_hint": "亿元"},
        "收入": {"max_value": 100000, "unit_hint": "亿元"},
        "研发投入": {"max_value": 10000, "unit_hint": "亿元"},
        "毛利率": {"min_value": -100, "max_value": 100, "unit_hint": "%"},
        "净利率": {"min_value": -100, "max_value": 100, "unit_hint": "%"},
        "增长率": {"min_value": -1000, "max_value": 1000, "unit_hint": "%"},
        "市占率": {"min_value": 0, "max_value": 100, "unit_hint": "%"},
        "份额": {"min_value": 0, "max_value": 100, "unit_hint": "%"},
        "销量": {"max_value": 10000, "unit_hint": "万辆"},
        "出口": {"max_value": 10000, "unit_hint": "万辆"},
    }

    @staticmethod
    def _validate_chart_datapoint(metric: str, value: float) -> bool:
        """校验单个数据点的指标-数值组合语义合理性"""
        for kw, rules in DocumentGenerationAgent._METRIC_SEMANTIC_RULES.items():
            if kw in metric:
                if "min_value" in rules and value < rules["min_value"]:
                    return False
                if "max_value" in rules and value > rules["max_value"]:
                    return False
        return True

    def _html_charts_from_datapoints(self, research_result: Dict[str, Any]) -> None:
        """从 data_points 结构化数据生成图表（最可靠）"""
        try:
            from src.services.chart_generator import ChartGenerator, ChartConfig, ChartType
        except ImportError:
            logger.warning("ChartGenerator not available for HTML charts")
            return
        except Exception:
            logger.exception("ChartGenerator import failed")
            return

        _root = Path(__file__).resolve().parent.parent.parent.parent
        chart_gen = ChartGenerator(output_dir=str(_root / "data" / "html_reports" / "charts"))

        sections = research_result.get("sections", [])
        for section in sections:
            data_points = section.get("data_points", [])
            if len(data_points) < 1:
                continue

            categories = []
            values = []
            for dp in data_points:
                metric = dp.get("metric", "")
                value = dp.get("value", "")
                unit = dp.get("unit", "")
                try:
                    v = float(re.sub(r'[^\d.\-]', '', str(value)))
                except (ValueError, TypeError):
                    continue
                # 语义校验：跳过明显不合理的数据点
                if not DocumentGenerationAgent._validate_chart_datapoint(metric, v):
                    logger.debug(f"Skipped semantically invalid datapoint: {metric}={v}{unit}")
                    continue
                categories.append(metric[:15])
                values.append(v)

            if len(categories) < 2:
                continue

            try:
                section_sources = section.get("sources", [])
                source_str = "Public data compilation"
                if section_sources and isinstance(section_sources, list):
                    titles = [s.get("title", "") for s in section_sources if isinstance(s, dict) and s.get("title")]
                    if titles:
                        source_str = "；".join(titles[:3])
                
                # P2: Detect chart type from data point metrics
                chart_type = ChartType.BAR
                metrics_lower = [dp.get("metric", "").lower() for dp in data_points]
                has_year = any("年" in m or "year" in m or "time" in m for m in metrics_lower)
                has_share = any("份额" in m or "share" in m or "占比" in m or "rate" in m for m in metrics_lower)
                
                if has_year and len(categories) >= 3:
                    chart_type = ChartType.LINE
                elif has_share and len(categories) <= 6:
                    chart_type = ChartType.PIE
                # else default BAR
                
                config = ChartConfig(
                    chart_type=chart_type,
                    title=f"{section.get('title', '')[:40]} - 关键数据",
                    data={"categories": categories[:10], "values": values[:10]},
                    xlabel="指标",
                    ylabel="数值",
                    source=source_str,
                )
                result = chart_gen.generate(config)
                if result.success and result.image_path:
                    existing = section.get("charts", []) or []
                    existing.append({
                        "path": result.image_path,
                        "caption": f"{section.get('title', '')} - 关键数据",
                        "section_title": section.get("title", ""),
                    })
                    section["charts"] = existing
                    logger.info(f"Generated data_points chart: {result.image_path}")
            except Exception:
                logger.exception(f"data_points chart failed: {section.get('title', '')}")

    def _html_charts_from_content(self, research_result: Dict[str, Any]) -> None:
        """使用 SmartChartGenerator 分析段落文本中的数据并配图"""
        try:
            from src.services.smart_chart_generator import SmartChartGenerator
        except ImportError:
            logger.warning("SmartChartGenerator not available for HTML charts")
            return

        _root = Path(__file__).resolve().parent.parent.parent.parent
        smart_gen = SmartChartGenerator(output_dir=str(_root / "data" / "html_reports" / "charts"))

        sections = research_result.get("sections", [])
        for section in sections:
            title = section.get("title", "")
            content = section.get("content", "")
            data_points = section.get("data_points", [])

            skip_keywords = ["概述", "总结", "结论", "建议", "方法论", "附录", "参考文献", "目录", "摘要"]
            if any(kw in title for kw in skip_keywords):
                continue

            try:
                suggestions = smart_gen.analyze_content(
                    section_title=title,
                    content=content,
                    data_points=data_points,
                )
                for suggestion in suggestions[:2]:
                    chart_path = smart_gen.generate_chart(suggestion)
                    if chart_path:
                        existing = section.get("charts", []) or []
                        existing.append({
                            "path": chart_path,
                            "caption": suggestion.caption or title,
                            "section_title": title,
                        })
                        section["charts"] = existing
                        logger.info(f"Generated content chart: {chart_path}")
                if section.get("charts"):
                    logger.info(f"Section '{title}': {len(section['charts'])} charts generated")
            except Exception:
                logger.exception(f"Content chart failed: {title}")

    # ==================== End chart fix ====================

    def _should_generate_charts(self, section_title: str) -> bool:
        """
        Determine whether charts should be generated for this section

        Args:
            section_title: Section title

        Returns:
            Whether charts should be generated
        """
        # Keywords suitable for chart generation (relaxed conditions)
        chart_keywords = [
            "市场规模", "市场份额", "竞争格局", "行业趋势",
            "财务分析", "用户分析", "技术对比", "区域分布",
            "增长分析", "投资分析", "销量分析", "数据对比",
            "占比", "排名", "趋势", "对比", "分布",
            "分析", "数据", "统计", "规模", "格局",  # New generic keywords
        ]
        
        # Keywords not suitable for chart generation
        no_chart_keywords = [
            "概述", "总结", "结论", "建议", "方法论",
            "附录", "参考文献", "目录", "摘要",
        ]
        
        # Check if contains unsuitable keywords
        for keyword in no_chart_keywords:
            if keyword in section_title:
                return False
        
        # Check if contains suitable keywords
        for keyword in chart_keywords:
            if keyword in section_title:
                return True
        
        return False
    
    def _fallback_generate_document(
        self,
        task_id: str,
        research_result: Dict[str, Any],
        output_format: str,
        output_path: Path
    ) -> Dict[str, Any]:
        """
        Fallback document generation method

        When DocumentGenerator is unavailable, use python-docx directly to generate document.

        Args:
            task_id: Task ID
            research_result: Research result
            output_format: Output format
            output_path: Output path

        Returns:
            Generation result dictionary
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = research_result.get("title", "Research Report")
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            if research_result.get("author"):
                doc.add_paragraph(f"Author: {research_result['author']}")
            if research_result.get("date"):
                doc.add_paragraph(f"Date: {research_result['date']}")
            
            doc.add_paragraph("")
            
# Sections
            for section in research_result.get("sections", []):
                section_title = section.get("title", "")
                section_content = section.get("content", "")
                
                if section_title:
                    doc.add_heading(section_title, level=1)
                
                if section_content:
                    for para in section_content.split("\n"):
                        if para.strip():
                            import re
                            img_match = re.match(r'^\s*<img\s[^>]*src="([^"]+)"[^>]*/?>\s*$', para)
                            if img_match:
                                try:
                                    doc.add_picture(img_match.group(1), width=Inches(5))
                                except Exception:
                                    doc.add_paragraph(para.strip())
                            else:
                                doc.add_paragraph(para.strip())
                
                # Subsections
                for subsection in section.get("subsections", []):
                    sub_title = subsection.get("title", "")
                    sub_content = subsection.get("content", "")
                    
                    if sub_title:
                        doc.add_heading(sub_title, level=2)
                    
                    if sub_content:
                        for para in sub_content.split("\n"):
                            if para.strip():
                                doc.add_paragraph(para.strip())
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            file_size = output_path.stat().st_size
            
            logger.info(f"Document generated via fallback: {output_path}")
            
            return DocumentGenerationResult(
                success=True,
                task_id=task_id,
                output_format=DocumentFormat.DOCX,
                version_id="v1",
                document_path=str(output_path),
                file_size=file_size,
                pages_estimate=max(1, len(research_result.get("sections", []))),
                message=f"文档已生成（备用方法）: {output_path.name}"
            ).to_dict()
            
        except ImportError:
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error="python-docx not installed, cannot generate document",
                error_code="DEPENDENCY_MISSING"
            ).to_dict()
        except Exception as e:
            logger.error(f"Fallback document generation failed: {e}", exc_info=True)
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error=f"Fallback generation failed: {str(e)}",
                error_code="FALLBACK_ERROR"
            ).to_dict()
    
    def _handle_regenerate_document(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle regeneration request

        Create new version based on historical research result

        Args:
            request: Document generation request

        Returns:
            Generation result (contains warning field indicating skeleton status)
        """
        return DocumentGenerationResult(
            success=True,
            task_id=request.task_id,
            output_format=request.output_format,
            version_id="v2",
            document_path=f"/output/{request.task_id}_report_v2.{request.output_format.value if request.output_format else 'docx'}",
            warning="SKELETON_IMPLEMENTATION: Regenerate logic pending (Week 22-23)"
        ).to_dict()
    
    async def _handle_adjust_content(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle content adjustment request

        Apply adjustments and create new version (revision loop support)

        Phase 8 refactoring: Uses SectionLocator and ContentApplier

        Args:
            request: Document generation request containing:
                - task_id: Task ID
                - section: Section name to adjust
                - adjustment: Adjustment description (user feedback)
                - document_path: Current document path
                - revision_type: Revision type (minor/section/phase/full)

        Returns:
            Adjustment result containing new document path
        """
        task_id = request.task_id or "unknown"
        
        # Extract parameters from adjustments (adjustments is List[Dict])
        adjustment_data = {}
        if request.adjustments and len(request.adjustments) > 0:
            adjustment_data = request.adjustments[0]  # 取第一个调整
        
        section = adjustment_data.get("section", "")
        section_id = adjustment_data.get("section_id")  # New: precise section ID
        adjustment = adjustment_data.get("adjustment", "")
        document_path = adjustment_data.get("document_path", "")
        revision_type = adjustment_data.get("revision_type", "minor")
        original_content = adjustment_data.get("original_content", "")
        keywords = adjustment_data.get("keywords", [])  # New: keyword search
        
        # Generate new version ID
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        version_id = f"v_{timestamp}"
        
        # === Phase 8: Use SectionLocator to locate section ===
        section_location = None
        if document_path and Path(document_path).exists():
            try:
                locator = SectionLocator()
                section_location = locator.locate(
                    document_path=document_path,
                    section_id=section_id,
                    section_title=section,
                    keywords=keywords,
                )
                
                if section_location:
                    logger.info(f"Located section: {section_location.section_id} - {section_location.section_title}")
                    # Get original content
                    if not original_content:
                        original_content = section_location.content
                else:
                    logger.warning(f"Section not found, using fallback: {section}")
                    
            except Exception as e:
                logger.warning(f"SectionLocator failed, using fallback: {e}")
        
        # Use LLM for content modification
        revised_content = None
        llm_success = False
        
        if adjustment:
            try:
                revision_prompt = self._build_revision_prompt(
                    section=section_location.section_title if section_location else section,
                    original_content=original_content,
                    adjustment=adjustment,
                    revision_type=revision_type,
                )
                
                llm_result = await call_llm(
                    prompt=revision_prompt,
                    system_prompt=self._get_revision_system_prompt(),
                    max_tokens=DEFAULT_LLM_MAX_TOKENS,
                    temperature=DEFAULT_LLM_TEMPERATURE,
                )
                
                if llm_result.get("success"):
                    revised_content = llm_result.get("content", "")
                    llm_success = True
                    logger.info(f"LLM revision successful for task {task_id}, section {section}")
                else:
                    logger.warning(f"LLM revision failed: {llm_result.get('error')}")
                    
            except Exception as e:
                logger.error(f"LLM revision error: {e}", exc_info=True)
        
        # === Phase 8: Use ContentApplier to apply revision ===
        revised_path = None
        apply_success = False
        
        if llm_success and revised_content and section_location and document_path:
            try:
                applier = ContentApplier(
                    backup_dir=str(self._documents_dir / "backups"),
                    create_backup=True,
                    version_suffix=True,
                )
                
                apply_result = applier.apply(
                    document_path=document_path,
                    location=section_location,
                    new_content=revised_content,
                    preserve_formatting=True,
                )
                
                if apply_result.success:
                    # Verify new document path exists
                    if apply_result.new_document_path and Path(apply_result.new_document_path).exists():
                        revised_path = Path(apply_result.new_document_path)
                        apply_success = True
                        logger.info(f"Content applied successfully: {revised_path}")
                    else:
                        logger.warning(f"ContentApplier succeeded but new document path invalid: {apply_result.new_document_path}")
                else:
                    logger.warning(f"ContentApplier failed: {apply_result.error}")
                    
            except Exception as e:
                logger.error(f"ContentApplier error: {e}", exc_info=True)
        
        # Fallback: Use old method to create revised version
        if not apply_success:
            if not llm_success:
                logger.info("Using legacy revision method: LLM revision not successful")
            elif not revised_content:
                logger.info("Using legacy revision method: No revised content generated")
            elif not section_location:
                logger.info("Using legacy revision method: Section not located")
            elif not document_path:
                logger.info("Using legacy revision method: No document path provided")
            else:
                logger.info("Using legacy revision method: ContentApplier failed")
            
            if document_path:
                try:
                    original_path = self._validate_document_path(document_path)
                except ValueError as e:
                    logger.warning(f"Invalid document path, using default: {e}")
                    original_path = self._documents_dir / Path(document_path).name
                
                filename = original_path.stem
                extension = original_path.suffix or ".docx"
                parent_dir = original_path.parent
                revised_path = parent_dir / f"{filename}_revised_{version_id}{extension}"
            else:
                revised_path = self._documents_dir / f"{task_id}_adjusted_{version_id}.docx"
            
            # If LLM modification successful, save revised content (old method)
            if llm_success and revised_content:
                try:
                    if document_path and Path(document_path).exists():
                        shutil.copy2(document_path, revised_path)
                    
                    # Save revised content to metadata file
                    meta_path = revised_path.with_suffix('.revision.json')
                    revision_meta = {
                        "task_id": task_id,
                        "version_id": version_id,
                        "section": section,
                        "adjustment": adjustment,
                        "revised_content": revised_content[:MAX_REVISION_CONTENT_LENGTH],
                        "created_at": datetime.now().isoformat(),
                        "llm_used": True,
                        "phase8_components_used": False,  # Mark Phase 8 components not used
                    }
                    
                    with open(meta_path, 'w', encoding='utf-8') as f:
                        json.dump(revision_meta, f, indent=2, ensure_ascii=False)
                        
                except OSError as e:
                    logger.error(f"Failed to save revised content: {e}")
        
        # Record revision to RevisionManager
        if self._revision_manager:
            try:
                revision_record = self._revision_manager.create_revision(
                    task_id=task_id,
                    revision_type=revision_type,
                    section=section,
                    adjustment=adjustment,
                    original_content=original_content[:MAX_REVISION_CONTENT_LENGTH] if original_content else None,
                    revised_content=revised_content[:MAX_REVISION_CONTENT_LENGTH] if revised_content else None,
                    document_path=document_path,
                    revised_document_path=str(revised_path) if revised_path else None,
                    user_feedback=adjustment,
                )
                version_id = revision_record.version_id
                logger.info(f"Revision recorded: {revision_record.revision_id}")
            except Exception as e:
                logger.warning(f"Failed to record revision: {e}")
        
        # Record revision info to SharedMemory (if available)
        if self._shared_memory:
            revision_record = {
                "task_id": task_id,
                "version_id": version_id,
                "section": section,
                "section_id": section_location.section_id if section_location else None,
                "adjustment": adjustment,
                "timestamp": datetime.now().isoformat(),
                "original_path": document_path,
                "revised_path": str(revised_path) if revised_path else None,
                "llm_used": llm_success,
                "phase8_components_used": apply_success,
            }
            self._shared_memory.set(f"revision_{task_id}_{version_id}", revision_record)
        
        # Send revision notification to MessageBus (if available)
        if self._message_bus:
            from src.core.communication import Event
            from src.core.orchestrator.execution.task_utils import safe_create_task
            safe_create_task(self._message_bus.publish(
                topic="document.adjusted",
                event=Event(
                    type="document.adjusted",
                    data={
                        "task_id": task_id,
                        "section": section,
                        "section_id": section_location.section_id if section_location else None,
                        "adjustment": adjustment,
                        "version_id": version_id,
                        "llm_used": llm_success,
                        "phase8_components_used": apply_success,
                    },
                    source=self.agent_id,
                )
            ), name="document_agent.publish_adjusted")
        
        # Build result message
        if apply_success:
            message = f"章节 '{section}' 已通过 Phase 8 组件修订完成"
        elif llm_success:
            message = f"章节 '{section}' 已通过LLM修订完成"
        else:
            message = f"章节 '{section}' 修订已记录（LLM未使用或失败）"
        
        return DocumentGenerationResult(
            success=True,
            task_id=task_id,
            output_format=request.output_format,
            version_id=version_id,
            document_path=str(revised_path) if revised_path else document_path,
            message=message,
            metadata={
                "llm_used": llm_success,
                "revision_type": revision_type,
                "section": section,
                "section_id": section_location.section_id if section_location else None,
                "phase8_components_used": apply_success,
            }
        ).to_dict()
    
    def _build_revision_prompt(
        self,
        section: str,
        original_content: str,
        adjustment: str,
        revision_type: str,
    ) -> str:
        """
        Build revision Prompt

        Args:
            section: Section name
            original_content: Original content
            adjustment: Adjustment description
            revision_type: Revision type

        Returns:
            Complete revision Prompt
        """
        type_instructions = {
            "minor": "请进行小幅调整，保持原有风格和结构。",
            "section": "请重写该章节，确保内容完整且符合要求。",
            "phase": "请重新分析该部分，提供更深入的见解。",
            "full": "请全面重写，确保所有内容都得到改进。",
        }
        
        instruction = type_instructions.get(revision_type, type_instructions["minor"])
        
        prompt = f"""你是一位专业的市场研究报告编辑。请根据用户的反馈修订以下章节内容。

## 章节名称
{section}

## 原始内容
{original_content[:MAX_PROMPT_CONTENT_LENGTH] if original_content else "（无原始内容）"}

## 用户反馈/修订要求
{adjustment}

## 修订指南
{instruction}

## 输出要求
1. 使用中文输出所有内容
2. 保持专业的研究报告风格
3. 确保数据准确性和来源可信
4. 使用清晰的结构和逻辑
5. 保持与原文档的一致性

请直接输出修订后的章节内容，不要包含解释或说明。
"""
        return prompt
    
    def _get_revision_system_prompt(self) -> str:
        """Get revision system Prompt"""
        return """你是一位资深的市场研究报告编辑，拥有国际咨询公司（McKinsey/BCG/Bain）从业背景。

## 专业领域
- 研究报告质量审核与修订
- 商业写作标准（金字塔原理、MECE原则）
- 数据准确性验证与来源追溯
- 逻辑论证链完整性检查

## 修订标准
1. **数据准确**：核实每个关键数据的来源和上下文，确保引用无误
2. **逻辑严谨**：每个结论必须有完整的论证链（判断→推导→数据→反面证据）
3. **表达精炼**：300字的有力论证 > 3000字的泛泛而谈
4. **结构清晰**：遵循金字塔原理，结论先行，层层展开

## 禁止行为
- 不得添加原文没有的新数据或事实
- 不得改变原文的核心结论
- 不得为了美观牺牲信息完整性

你的输出应直接是修订后的内容，不要包含解释说明。"""

    def _handle_export_document(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle export request

        Copy document to user-specified path

        Args:
            request: Document generation request

        Returns:
            Export result (contains warning field indicating skeleton status)
        """
        return DocumentGenerationResult(
            success=True,
            task_id=request.task_id,
            output_format=request.output_format,
            document_path=request.export_path,
            warning="SKELETON_IMPLEMENTATION: Export logic pending (Week 24)"
        ).to_dict()
    
    def _handle_get_preview(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle preview request - Generate HTML preview

        Flow: research_result → ContentOrchestrator → HTML → Return HTML path

        Args:
            request: Document generation request

        Returns:
            Preview result (containing HTML path)
        """
        task_id = request.task_id or f"preview_{uuid.uuid4().hex[:8]}"
        
        # 1. Get research content
        research_result = request.research_result
        if not research_result:
            research_result = self._load_research_result(task_id)
        
        if not research_result:
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error="No research result for preview",
                error_code="NO_CONTENT"
            ).to_dict()
        
        # 2. Generate charts for HTML before creating preview
        research_result = self._generate_charts_for_html(research_result)
        
        # 3. Use ContentOrchestrator to generate HTML preview
        try:
            from src.content.content_orchestrator import ContentOrchestrator
            
            from src.core.preview_storage import PreviewStorage
            
            orchestrator = ContentOrchestrator()
            html_content = orchestrator.transform_to_html(
                research_result=research_result,
                output_format="html",  # Fix: Use "html" to enable chart rendering (is_html_format=True)
                output_dir=str(PreviewStorage.NEW_DIR)  # External chart images
            )
            
            # 3. Save HTML file to serving directory
            PreviewStorage.write(task_id, html_content)
            
            logger.info(f"Generated HTML preview for {task_id}")
            
            return DocumentGenerationResult(
                success=True,
                task_id=task_id,
                output_format=request.output_format,
                document_path=str(html_path),  # Fix: Use document_path for orchestrator compatibility
                preview_path=str(html_path)
            ).to_dict()
            
        except Exception as e:
            logger.error(f"Failed to generate HTML preview: {e}")
            return DocumentGenerationResult(
                success=False,
                task_id=task_id,
                error=f"Preview generation failed: {e}",
                error_code="PREVIEW_FAILED"
            ).to_dict()
    
    def _handle_list_versions(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle version list request

        Args:
            request: Document generation request

        Returns:
            Version list
        """
        # Week 21 skeleton implementation
        return DocumentGenerationResult(
            success=True,
            task_id=request.task_id,
            output_format=request.output_format,
            versions=[
                {
                    "version_id": "v1",
                    "created_at": datetime.now().isoformat(),
                    "created_by": "initial"
                }
            ]
        ).to_dict()
    
    def _handle_rollback_version(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle version rollback request

        Args:
            request: Document generation request

        Returns:
            Comparison result
        """
        # Week 21 skeleton implementation
        return DocumentGenerationResult(
            success=True,
            task_id=request.task_id,
            output_format=request.output_format,
            version_id=f"v_rollback_from_{request.version_id}"
        ).to_dict()
    
    def _handle_compare_versions(
        self,
        request: DocumentGenerationRequest
    ) -> Dict[str, Any]:
        """
        Handle version comparison request

        Args:
            request: Document generation request

        Returns:
            Comparison result
        """
        # Week 21 skeleton implementation
        return DocumentGenerationResult(
            success=True,
            task_id=request.task_id,
            output_format=request.output_format,
            diff_result={
                "version_1": request.version_id,
                "version_2": request.version_id_2,
                "changes": []  # To be implemented
            }
        ).to_dict()


# Export
__all__ = ["DocumentGenerationAgent"]