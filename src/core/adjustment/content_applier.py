# -*- coding: utf-8 -*-
"""
ContentApplier - 内容应用器

Phase 8: 报告修订闭环

将 LLM 生成的新内容应用到文档指定位置。

设计文档: docs/KNOWLEDGE_BASE/02_ARCHITECTURE/PHASE8_DEVELOPMENT_PLAN.md
"""

import html
import logging
import re
import shutil
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

from .section_locator import SectionLocation

logger = logging.getLogger(__name__)


@dataclass
class ApplyResult:
    """
    内容应用结果
    
    Attributes:
        success: 是否成功
        new_document_path: 新文档路径
        backup_path: 备份路径
        changes: 变更详情
        error: 错误信息
    """
    success: bool
    new_document_path: Optional[str] = None
    backup_path: Optional[str] = None
    changes: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "success": self.success,
            "new_document_path": self.new_document_path,
            "backup_path": self.backup_path,
            "changes": self.changes,
            "error": self.error,
        }


class ContentApplier:
    """
    内容应用器
    
    将修订内容应用到文档指定位置，支持备份和版本管理。
    
    使用示例:
        applier = ContentApplier(backup_dir="data/backups")
        
        result = applier.apply(
            document_path="/path/to/report.docx",
            location=section_location,
            new_content="新的章节内容...",
        )
        
        if result.success:
            print(f"Updated: {result.new_document_path}")
    """
    
    def __init__(
        self,
        backup_dir: str = "data/backups",
        create_backup: bool = True,
        version_suffix: bool = True,
    ):
        """
        初始化内容应用器
        
        Args:
            backup_dir: 备份目录
            create_backup: 是否创建备份
            version_suffix: 是否在文件名中添加版本后缀
        """
        self._backup_dir = Path(backup_dir)
        self._create_backup = create_backup
        self._version_suffix = version_suffix
        
        # 创建备份目录
        if self._create_backup and not self._backup_dir.exists():
            self._backup_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created backup directory: {self._backup_dir}")
        
        logger.info("ContentApplier initialized")
    
    def apply(
        self,
        document_path: str,
        location: SectionLocation,
        new_content: str,
        preserve_formatting: bool = True,
    ) -> ApplyResult:
        """
        应用修订内容
        
        Args:
            document_path: 文档路径
            location: 章节位置信息
            new_content: 新内容
            preserve_formatting: 是否保留原有格式
            
        Returns:
            ApplyResult
        """
        logger.info(f"Applying content to section: {location.section_id}")
        
        # 1. 验证输入
        if not Path(document_path).exists():
            return ApplyResult(
                success=False,
                error=f"Document not found: {document_path}",
            )
        
        if not new_content or not new_content.strip():
            return ApplyResult(
                success=False,
                error="New content is empty",
            )
        
        # 2. 创建备份
        backup_path = None
        if self._create_backup:
            backup_path = self._create_backup_file(document_path)
        
        # 3. 根据文档类型选择应用方式
        suffix = Path(document_path).suffix.lower()
        
        try:
            if suffix == '.docx':
                new_path = self._apply_to_docx(
                    document_path, location, new_content, preserve_formatting
                )
            elif suffix == '.html':
                new_path = self._apply_to_html(
                    document_path, location, new_content
                )
            elif suffix == '.md':
                new_path = self._apply_to_markdown(
                    document_path, location, new_content
                )
            else:
                return ApplyResult(
                    success=False,
                    error=f"Unsupported format: {suffix}",
                )
            
            # 4. 计算变更
            changes = {
                "section_id": location.section_id,
                "section_title": location.section_title,
                "old_length": len(location.content),
                "new_length": len(new_content),
                "timestamp": datetime.now().isoformat(),
            }
            
            logger.info(f"Content applied successfully: {new_path}")
            
            return ApplyResult(
                success=True,
                new_document_path=str(new_path),
                backup_path=str(backup_path) if backup_path else None,
                changes=changes,
            )
            
        except Exception as e:
            logger.error(f"Failed to apply content: {e}")
            return ApplyResult(
                success=False,
                error=str(e),
                backup_path=str(backup_path) if backup_path else None,
            )
    
    def apply_global_style(
        self,
        document_path: str,
        style_changes: Dict[str, Any],
    ) -> ApplyResult:
        """
        应用全局样式变更
        
        Args:
            document_path: 文档路径
            style_changes: 样式变更（字体、颜色、边距等）
            
        Returns:
            ApplyResult
        """
        suffix = Path(document_path).suffix.lower()
        
        # 创建备份
        backup_path = None
        if self._create_backup:
            backup_path = self._create_backup_file(document_path)
        
        try:
            if suffix == '.docx':
                new_path = self._apply_style_to_docx(document_path, style_changes)
            else:
                return ApplyResult(
                    success=False,
                    error=f"Style changes not supported for: {suffix}",
                )
            
            return ApplyResult(
                success=True,
                new_document_path=str(new_path),
                backup_path=str(backup_path) if backup_path else None,
                changes={"style_changes": style_changes},
            )
            
        except Exception as e:
            return ApplyResult(
                success=False,
                error=str(e),
            )
    
    # ==================== 内部方法 ====================
    
    def _create_backup_file(self, document_path: str) -> Path:
        """创建备份文件"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = Path(document_path)
        
        backup_name = f"{path.stem}_backup_{timestamp}{path.suffix}"
        backup_path = self._backup_dir / backup_name
        
        shutil.copy2(document_path, backup_path)
        logger.debug(f"Created backup: {backup_path}")
        
        return backup_path
    
    def _apply_to_docx(
        self,
        document_path: str,
        location: SectionLocation,
        new_content: str,
        preserve_formatting: bool = True,
    ) -> Path:
        """
        应用到 Word 文档
        
        Args:
            document_path: 文档路径
            location: 章节位置
            new_content: 新内容
            preserve_formatting: 是否保留格式
            
        Returns:
            新文档路径
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word document editing")
        
        document = Document(document_path)
        
        # 查找目标章节段落
        target_paragraphs = self._find_section_paragraphs(
            document, location.section_title
        )
        
        if not target_paragraphs:
            raise ValueError(f"Section not found: {location.section_title}")
        
        # 替换内容
        if preserve_formatting and target_paragraphs:
            # 保留第一个段落格式，替换内容
            first_para = target_paragraphs[0]
            first_para.text = new_content
            
            # 删除多余段落
            for para in target_paragraphs[1:]:
                p = para._element
                p.getparent().remove(p)
        else:
            # 简单替换所有段落
            for i, para in enumerate(target_paragraphs):
                if i == 0:
                    para.text = new_content
                else:
                    p = para._element
                    p.getparent().remove(p)
        
        # 保存新文档
        new_path = self._get_new_path(document_path)
        document.save(str(new_path))
        
        return new_path
    
    def _find_section_paragraphs(
        self,
        document,
        section_title: str,
    ) -> list:
        """
        查找章节相关的段落
        
        Args:
            document: Document 对象
            section_title: 章节标题
            
        Returns:
            段落列表
        """
        paragraphs = list(document.paragraphs)
        target_paragraphs = []
        
        # 查找标题段落
        title_index = -1
        for i, para in enumerate(paragraphs):
            if section_title.lower() in para.text.lower():
                style_name = para.style.name.lower() if para.style else ""
                if any(h in style_name for h in ['heading', '标题', 'title']):
                    title_index = i
                    target_paragraphs.append(para)
                    break
        
        if title_index == -1:
            return []
        
        # 收集章节内容段落（直到下一个标题）
        title_style = paragraphs[title_index].style
        
        for i in range(title_index + 1, len(paragraphs)):
            para = paragraphs[i]
            style_name = para.style.name.lower() if para.style else ""
            
            # 遇到下一个标题则停止
            if any(h in style_name for h in ['heading', '标题']):
                break
            
            target_paragraphs.append(para)
        
        return target_paragraphs
    
    def _apply_style_to_docx(
        self,
        document_path: str,
        style_changes: Dict[str, Any],
    ) -> Path:
        """应用样式到 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx is required")
        
        document = Document(document_path)
        
        # 应用字体变更
        if "font_name" in style_changes or "font_size" in style_changes:
            for para in document.paragraphs:
                for run in para.runs:
                    if "font_name" in style_changes:
                        run.font.name = style_changes["font_name"]
                    if "font_size" in style_changes:
                        run.font.size = Pt(style_changes["font_size"])
        
        # 保存
        new_path = self._get_new_path(document_path)
        document.save(str(new_path))
        
        return new_path
    
    def _apply_to_html(
        self,
        document_path: str,
        location: SectionLocation,
        new_content: str,
    ) -> Path:
        """应用到 HTML 文档"""
        with open(document_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 查找章节位置
        # 简化处理：替换标题后的内容直到下一个标题
        import re
        
        # 查找标题
        title_pattern = rf'(<h[1-6][^>]*>{re.escape(location.section_title)}</h\1>)'
        match = re.search(title_pattern, content, re.IGNORECASE)
        
        if not match:
            raise ValueError(f"Section title not found: {location.section_title}")
        
        # 找到下一个标题位置
        title_end = match.end()
        next_heading = re.search(r'<h[1-6][^>]*>', content[title_end:], re.IGNORECASE)
        
        if next_heading:
            section_end = title_end + next_heading.start()
        else:
            # 没有下一个标题，到文档结束
            section_end = len(content)
        
        # 替换内容 (XSS protection: escape content)
        old_section = content[title_end:section_end]
        safe_content = html.escape(new_content)
        new_content_wrapped = f"\n<p>{safe_content}</p>\n"
        
        new_document = content[:title_end] + new_content_wrapped + content[section_end:]
        
        # 保存
        new_path = self._get_new_path(document_path)
        with open(new_path, 'w', encoding='utf-8') as f:
            f.write(new_document)
        
        return new_path
    
    def _apply_to_markdown(
        self,
        document_path: str,
        location: SectionLocation,
        new_content: str,
    ) -> Path:
        """应用到 Markdown 文档"""
        with open(document_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 查找标题行
        title_pattern = rf'^#+\s+{re.escape(location.section_title)}\s*$'
        
        title_line = -1
        for i, line in enumerate(lines):
            if re.match(title_pattern, line, re.IGNORECASE):
                title_line = i
                break
        
        if title_line == -1:
            raise ValueError(f"Section title not found: {location.section_title}")
        
        # 找到下一个标题
        next_title = -1
        for i in range(title_line + 1, len(lines)):
            if re.match(r'^#+\s+', lines[i]):
                next_title = i
                break
        
        if next_title == -1:
            next_title = len(lines)
        
        # 替换内容
        new_lines = lines[:title_line + 1] + [new_content + "\n"] + lines[next_title:]
        
        # 保存
        new_path = self._get_new_path(document_path)
        with open(new_path, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
        
        return new_path
    
    def _get_new_path(self, document_path: str) -> Path:
        """生成新文件路径"""
        path = Path(document_path)
        
        if self._version_suffix:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_name = f"{path.stem}_v{timestamp}{path.suffix}"
        else:
            new_name = path.name
        
        return path.parent / new_name

    # ==================== Section Insertion (Phase 8: Revision) ====================

    def insert_section(
        self,
        document_path: str,
        new_content: str,
        section_title: str,
        location: Optional[SectionLocation] = None,
        level: int = 2,
    ) -> ApplyResult:
        """
        Insert a new section into the document.

        Args:
            document_path: Path to the document
            new_content: HTML content of the new section
            section_title: Title of the new section
            location: Insertion point (None = append at end)
            level: Heading level (2 = h2)

        Returns:
            ApplyResult with new_document_path
        """
        backup_path = self._create_backup_file(document_path) if self._create_backup else None

        suffix = Path(document_path).suffix.lower()
        if suffix == '.html':
            return self._insert_section_into_html(
                document_path, new_content, section_title, location, level, backup_path
            )
        return ApplyResult(
            success=False,
            error=f"Section insertion not supported for format: {suffix}",
        )

    def _insert_section_into_html(
        self,
        document_path: str,
        new_content: str,
        section_title: str,
        location: Optional[SectionLocation],
        level: int,
        backup_path: Optional[Path],
    ) -> ApplyResult:
        """Insert a new section into an HTML document (XSS protected)."""
        with open(document_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # XSS protection: escape all user/LLM generated content
        safe_title = html.escape(section_title)
        safe_content = html.escape(new_content)

        # Generate valid HTML ID (slug: lowercase, no spaces, no special chars)
        section_id = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff_-]', '', section_title.strip())
        section_id = section_id.replace(' ', '-').lower()

        heading_tag = f"h{level}"
        section_html = f'\n<{heading_tag} id="{section_id}">{safe_title}</{heading_tag}>\n<p>{safe_content}</p>\n'

        if location and location.end_index > 0:
            # Insert after the specified section
            insert_pos = location.end_index
            html_content = html_content[:insert_pos] + section_html + html_content[insert_pos:]
        else:
            # Append before </body>
            if '</body>' in html_content:
                html_content = html_content.replace('</body>', f'{section_html}\n</body>')
            else:
                html_content += section_html

        with open(document_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return ApplyResult(
            success=True,
            new_document_path=document_path,
            backup_path=str(backup_path) if backup_path else None,
            changes={"action": "insert_section", "title": safe_title, "level": level},
        )

    def rebuild_toc(self, document_path: str) -> bool:
        """
        Rebuild the Table of Contents in an HTML document.

        Scans all <h1>-<h6> tags, generates anchor IDs, and updates
        the TOC navigation block. Also injects id attributes into
        heading tags for navigation.

        Args:
            document_path: Path to the HTML document

        Returns:
            bool: True if TOC was rebuilt successfully
        """
        with open(document_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # Extract all headings with three capture groups: level, attrs, title text
        heading_pattern = r'<h([1-6])([^>]*)>(.*?)</h\1>'
        headings = re.findall(heading_pattern, html_content, re.DOTALL)

        # Build TOC items and track used IDs to prevent collisions
        used_ids: Set[str] = set()
        toc_items: List[str] = []

        for level_num, existing_attrs, title_text in headings:
            indent = "  " * (int(level_num) - 1)
            safe_text = html.escape(title_text.strip())

            # Generate unique anchor ID
            base_id = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff_-]', '', title_text.strip())
            base_id = base_id.replace(' ', '-').lower()
            section_id = base_id
            counter = 1
            while section_id in used_ids:
                section_id = f"{base_id}_{counter}"
                counter += 1
            used_ids.add(section_id)

            toc_items.append(f'{indent}<li><a href="#{section_id}">{safe_text}</a></li>')

        toc_html = '<nav id="toc">\n<ul>\n' + '\n'.join(toc_items) + '\n</ul>\n</nav>'

        # Inject id attributes into heading tags (preserving existing attributes)
        used_ids_for_injection: Set[str] = set()

        def _add_id_to_heading(match: re.Match) -> str:
            level = match.group(1)
            existing_attrs = match.group(2) or ""
            title_text = match.group(3)

            # Check if id already exists
            if 'id=' in existing_attrs.lower():
                return match.group(0)

            # Generate unique ID
            base_id = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff_-]', '', title_text.strip())
            base_id = base_id.replace(' ', '-').lower()
            section_id = base_id
            counter = 1
            while section_id in used_ids_for_injection:
                section_id = f"{base_id}_{counter}"
                counter += 1
            used_ids_for_injection.add(section_id)

            return f'<h{level}{existing_attrs} id="{section_id}">{title_text}</h{level}>'

        html_content = re.sub(
            heading_pattern,
            _add_id_to_heading,
            html_content,
            flags=re.DOTALL,
        )

        # Replace existing TOC or inject new one
        if '<nav id="toc">' in html_content:
            html_content = re.sub(
                r'<nav id="toc">.*?</nav>',
                toc_html,
                html_content,
                flags=re.DOTALL,
            )
        else:
            html_content = html_content.replace('<body>', f'<body>\n{toc_html}', 1)

        with open(document_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return True


__all__ = ["ContentApplier", "ApplyResult"]
