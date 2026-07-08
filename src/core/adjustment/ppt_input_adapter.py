from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from src.core.adjustment.extraction_types import DataParser, ExtractionResult
from src.content.content_orchestrator import ContentSection, SectionType


class DocxDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        import docx
        doc = docx.Document(file_path)

        sections: List[ContentSection] = []
        current_section: Optional[ContentSection] = None
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                current_section = ContentSection(
                    id=f"sec_{len(sections)}",
                    title=para.text.strip(),
                    content="",
                    order=len(sections),
                    type=SectionType.BODY,
                    points=[],
                )
                sections.append(current_section)
            elif current_section and para.text.strip():
                current_section.points.append(para.text.strip())

        if not sections:
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if text:
                sections = [ContentSection(
                    id="sec_0", title="", content=text,
                    order=0, type=SectionType.BODY, points=[],
                )]

        tables: List[List[List[str]]] = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)

        title = self._detect_title(doc, sections)
        key_topics = self._extract_topics(sections)

        return ExtractionResult(
            title=title,
            sections=sections,
            tables=tables,
            key_topics=key_topics,
            metadata={
                "format": "docx",
                "para_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
            summary=None,
        )

    def _detect_title(self, doc, sections: List[ContentSection]) -> str:
        if sections and sections[0].title:
            return sections[0].title
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading 1") and para.text.strip():
                return para.text.strip()
        return ""

    def _extract_topics(self, sections: List[ContentSection]) -> List[str]:
        return [s.title for s in sections if s.title]


class PdfDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        sections: List[ContentSection] = []
        tables: List[List[List[str]]] = []
        total_pages = 0

        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        sections.append(ContentSection(
                            id=f"sec_{i}",
                            title=f"Page {i + 1}",
                            content=text.strip(),
                            order=i,
                            type=SectionType.BODY,
                            points=[],
                        ))
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            tables.append(table)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        sections.append(ContentSection(
                            id=f"sec_{i}",
                            title=f"Page {i + 1}",
                            content=text.strip(),
                            order=i,
                            type=SectionType.BODY,
                            points=[],
                        ))
            except ImportError:
                pass

        key_topics = self._extract_topics(sections)

        return ExtractionResult(
            title=sections[0].title if sections else "",
            sections=sections,
            tables=tables,
            key_topics=key_topics,
            metadata={"format": "pdf", "page_count": total_pages},
            summary=None,
        )

    def _extract_topics(self, sections: List[ContentSection]) -> List[str]:
        return [s.title for s in sections if s.title and not s.title.startswith("Page ")]


class ExcelDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        sections: List[ContentSection] = []
        tables: List[List[List[str]]] = []
        sheet_count = len(wb.sheetnames)

        for idx, ws_name in enumerate(wb.sheetnames):
            ws = wb[ws_name]
            rows: List[List[str]] = []
            for row in ws.iter_rows(values_only=True):
                str_row = [str(cell) if cell is not None else "" for cell in row]
                rows.append(str_row)

            if rows:
                tables.append(rows)
                content = "\n".join(" | ".join(r) for r in rows[:20])
                sections.append(ContentSection(
                    id=f"sec_{idx}",
                    title=ws_name,
                    content=content,
                    order=idx,
                    type=SectionType.BODY,
                    points=[],
                ))

        wb.close()

        return ExtractionResult(
            title=sections[0].title if sections else p.stem,
            sections=sections,
            tables=tables,
            key_topics=[s.title for s in sections if s.title],
            metadata={"format": "xlsx", "sheet_count": sheet_count},
            summary=None,
        )


class TextDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        text = p.read_text(encoding="utf-8", errors="replace").strip()
        sections: List[ContentSection] = []
        if text:
            sections.append(ContentSection(
                id="sec_0",
                title=p.stem,
                content=text,
                order=0,
                type=SectionType.BODY,
                points=[],
            ))

        return ExtractionResult(
            title=p.stem,
            sections=sections,
            tables=[],
            key_topics=[],
            metadata={"format": "txt", "word_count": len(text.split())},
            summary=None,
        )


class CsvDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        import csv
        rows: List[List[str]] = []
        with open(file_path, "r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)

        sections: List[ContentSection] = []
        if rows:
            content = "\n".join(" | ".join(r) for r in rows[:50])
            sections.append(ContentSection(
                id="sec_0",
                title=p.stem,
                content=content,
                order=0,
                type=SectionType.BODY,
                points=[],
            ))

        return ExtractionResult(
            title=p.stem,
            sections=sections,
            tables=[rows] if rows else [],
            key_topics=rows[0] if rows else [],
            metadata={"format": "csv", "row_count": len(rows)},
            summary=None,
        )


class JsonDataParser(DataParser):
    def parse(self, file_path: str) -> ExtractionResult:
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        import json
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        sections: List[ContentSection] = []
        tables: List[List[List[str]]] = []
        title = p.stem

        if isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, dict):
                    content = "\n".join(f"{k}: {v}" for k, v in item.items())
                    sections.append(ContentSection(
                        id=f"sec_{i}",
                        title=item.get("name", item.get("heading", f"Item {i + 1}")),
                        content=content,
                        order=i,
                        type=SectionType.BODY,
                        points=[],
                    ))
                else:
                    sections.append(ContentSection(
                        id=f"sec_{i}",
                        title=f"Item {i + 1}",
                        content=str(item),
                        order=i,
                        type=SectionType.BODY,
                        points=[],
                    ))
        elif isinstance(data, dict):
            title = data.get("title", p.stem)
            if "sections" in data and isinstance(data["sections"], list):
                for i, sec in enumerate(data["sections"]):
                    if isinstance(sec, dict):
                        sections.append(ContentSection(
                            id=f"sec_{i}",
                            title=sec.get("heading", sec.get("title", f"Section {i + 1}")),
                            content=sec.get("text", str(sec)),
                            order=i,
                            type=SectionType.BODY,
                            points=[],
                        ))
            else:
                content = "\n".join(f"{k}: {v}" for k, v in data.items())
                sections.append(ContentSection(
                    id="sec_0",
                    title=title,
                    content=content,
                    order=0,
                    type=SectionType.BODY,
                    points=[],
                ))

        return ExtractionResult(
            title=title,
            sections=sections,
            tables=tables,
            key_topics=[s.title for s in sections if s.title],
            metadata={"format": "json"},
            summary=None,
        )


class PptInputAdapter:
    def __init__(self):
        self._parsers: Dict[str, DataParser] = {
            ".docx": DocxDataParser(),
            ".pdf": PdfDataParser(),
            ".xlsx": ExcelDataParser(),
            ".xls": ExcelDataParser(),
            ".txt": TextDataParser(),
            ".md": TextDataParser(),
            ".csv": CsvDataParser(),
            ".json": JsonDataParser(),
        }

    def extract(self, file_paths: List[str]) -> ExtractionResult:
        results: List[ExtractionResult] = []
        for fp in file_paths:
            ext = Path(fp).suffix.lower()
            parser = self._parsers.get(ext)
            if parser:
                results.append(parser.parse(fp))
            else:
                results.append(self._fallback_parse(fp))
        return self._merge(results)

    def _fallback_parse(self, file_path: str) -> ExtractionResult:
        return ExtractionResult(
            title=Path(file_path).stem,
            sections=[],
            tables=[],
            key_topics=[],
            metadata={"format": "unknown", "path": file_path},
            summary=None,
        )

    def _merge(self, results: List[ExtractionResult]) -> ExtractionResult:
        if not results:
            return ExtractionResult(
                title="", sections=[], tables=[],
                key_topics=[], metadata={}, summary=None,
            )
        if len(results) == 1:
            return results[0]

        all_sections: List[ContentSection] = []
        all_tables: List[List[List[str]]] = []
        all_topics: List[str] = []
        merged_meta: Dict[str, Any] = {"formats": [], "file_count": len(results)}
        for r in results:
            all_sections.extend(r.sections)
            all_tables.extend(r.tables)
            all_topics.extend(r.key_topics)
            if r.metadata.get("format"):
                merged_meta["formats"].append(r.metadata["format"])

        title = results[0].title if results[0].title else ""

        return ExtractionResult(
            title=title,
            sections=all_sections,
            tables=all_tables,
            key_topics=all_topics,
            metadata=merged_meta,
            summary=None,
        )
