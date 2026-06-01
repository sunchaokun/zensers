# -*- coding: utf-8 -*-
"""
Document Preview Generator
========================

Provides document preview generation:
1. Generate preview images (PNG)
2. Preview cache management
3. Multi-format support

Technical approach:
- Word/PPT: python-docx/python-pptx → basic preview
- PDF: pdf2image → PNG (optional dependency)
"""

import hashlib
import json
import logging
import os
import re
import tempfile
import threading
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# Constants
VALID_FORMATS = ['docx', 'pptx', 'pdf', 'html']
DEFAULT_PREVIEW_WIDTH = 1024
DEFAULT_PREVIEW_HEIGHT = 768
MAX_PREVIEW_SIZE = 10 * 1024 * 1024  # 10MB


@dataclass
class PreviewResult:
    """Preview result"""
    success: bool
    preview_path: Optional[str] = None
    preview_format: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None
    file_size: Optional[int] = None
    cached: bool = False
    error: Optional[str] = None
    error_code: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "success": self.success,
            "preview_path": self.preview_path,
            "preview_format": self.preview_format,
            "width": self.width,
            "height": self.height,
            "file_size": self.file_size,
            "cached": self.cached,
            "error": self.error,
            "error_code": self.error_code
        }


class PreviewGenerator:
    """
    Document Preview Generator
    
    Generates preview images for documents with caching support.
    
    Usage example:
        generator = PreviewGenerator(cache_dir="data/previews")
        
        result = generator.generate_preview(
            document_path="/path/to/report.docx",
            format="docx"
        )
        
        if result.success:
            print(f"Preview: {result.preview_path}")
    """
    
    def __init__(self, cache_dir: str = "data/previews"):
        """
        Initialize preview generator
        
        Args:
            cache_dir: Preview cache directory
        """
        self.cache_dir = Path(cache_dir)
        
        # Thread lock to protect concurrent access to cache index
        self._lock = threading.Lock()
        
        # Create cache directory
        if not self.cache_dir.exists():
            self.cache_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created preview cache directory: {self.cache_dir}")
        
        # Preview cache index
        self._cache_index: Dict[str, str] = {}
        self._load_cache_index()
        
        # Check optional dependencies
        self._pdf2image_available = self._check_pdf2image()
    
    def _check_pdf2image(self) -> bool:
        """Check if pdf2image is available"""
        try:
            from pdf2image import convert_from_path  # noqa: F401
            return True
        except ImportError:
            return False
    
    def _load_cache_index(self) -> None:
        """Load cache index (thread-safe)"""
        import json
        index_file = self.cache_dir / "cache_index.json"
        
        with self._lock:
            if index_file.exists():
                try:
                    with open(index_file, 'r', encoding='utf-8') as f:
                        self._cache_index = json.load(f)
                except (json.JSONDecodeError, IOError) as e:
                    logger.warning(f"Failed to load cache index: {e}")
                    self._cache_index = {}
    
    def _save_cache_index(self) -> None:
        """Save cache index (thread-safe)"""
        index_file = self.cache_dir / "cache_index.json"
        
        with self._lock:
            try:
                import json
                with open(index_file, 'w', encoding='utf-8') as f:
                    json.dump(self._cache_index, f, ensure_ascii=False, indent=2)
            except IOError as e:
                logger.warning(f"Failed to save cache index: {e}")
    
    def _get_cache_key(self, document_path: str, **options) -> str:
        """Generate cache key"""
        # Based on file path + modification time + options
        mtime = os.path.getmtime(document_path) if os.path.exists(document_path) else 0
        key_data = f"{document_path}:{mtime}:{sorted(options.items())}"
        return hashlib.md5(key_data.encode(), usedforsecurity=False).hexdigest()
    
    def _is_valid_format(self, format: str) -> bool:
        """Validate format"""
        return format in VALID_FORMATS
    
    def _generate_placeholder_result(self, document_path: str, format: str) -> PreviewResult:
        """Generate placeholder preview for oversized files"""
        # Generate placeholder image
        placeholder_path = self._generate_placeholder_preview(
            document_path,
            width=DEFAULT_PREVIEW_WIDTH,
            height=DEFAULT_PREVIEW_HEIGHT,
        )
        
        if placeholder_path and os.path.exists(placeholder_path):
            return PreviewResult(
                success=True,
                preview_path=placeholder_path,
                preview_format="png",
                width=DEFAULT_PREVIEW_WIDTH,
                height=DEFAULT_PREVIEW_HEIGHT,
                file_size=os.path.getsize(placeholder_path),
                cached=False,
                error="Document too large, showing placeholder"
            )
        
        return PreviewResult(
            success=False,
            error="Document too large to preview",
            error_code="DOCUMENT_TOO_LARGE"
        )
    
    def generate_preview(
        self,
        document_path: str,
        format: str,
        width: int = DEFAULT_PREVIEW_WIDTH,
        height: int = DEFAULT_PREVIEW_HEIGHT,
        page_number: int = 1,
        use_cache: bool = True
    ) -> PreviewResult:
        """
        Generate preview
        
        Args:
            document_path: Document path
            format: Document format
            width: Preview width
            height: Preview height
            page_number: Page number (1-based)
            use_cache: Whether to use cache
            
        Returns:
            PreviewResult preview result
        """
        # Validate format
        if not self._is_valid_format(format):
            logger.warning(f"Invalid format: {format}")
            return PreviewResult(
                success=False,
                error=f"Invalid format: {format}",
                error_code="INVALID_FORMAT"
            )
        
        # Validate file exists
        if not os.path.exists(document_path):
            logger.warning(f"Document not found: {document_path}")
            return PreviewResult(
                success=False,
                error="Document not found",
                error_code="DOCUMENT_NOT_FOUND"
            )
        
        # P0-4 fix: check if file size exceeds limit
        file_size = os.path.getsize(document_path)
        if file_size > MAX_PREVIEW_SIZE:
            logger.warning(f"Document size ({file_size} bytes) exceeds MAX_PREVIEW_SIZE ({MAX_PREVIEW_SIZE} bytes), using placeholder")
            # For oversized files, return placeholder preview
            return self._generate_placeholder_result(document_path, format)
        
        # Check cache (thread-safe)
        cache_key = None
        if use_cache:
            cache_key = self._get_cache_key(document_path, format=format, width=width, height=height, page=page_number)
            
            with self._lock:
                if cache_key in self._cache_index:
                    cached_path = self._cache_index[cache_key]
                    if os.path.exists(cached_path):
                        logger.info(f"Using cached preview: {cached_path}")
                        return PreviewResult(
                            success=True,
                            preview_path=cached_path,
                            preview_format="png",
                            cached=True,
                            file_size=os.path.getsize(cached_path)
                        )
        
        try:
            # Generate preview by format
            if format == "pdf":
                preview_path = self._generate_pdf_preview(document_path, width, height, page_number)
            elif format in ["docx", "pptx"]:
                preview_path = self._generate_office_preview(document_path, format, width, height, page_number)
            else:
                preview_path = self._generate_html_preview(document_path, width, height)
            
            if preview_path and os.path.exists(preview_path):
                file_size = os.path.getsize(preview_path)
                
                # Update cache (thread-safe)
                if use_cache and cache_key:
                    with self._lock:
                        self._cache_index[cache_key] = preview_path
                    self._save_cache_index()
                
                logger.info(f"Generated preview: {preview_path}")
                
                # Return appropriate preview format
                if format == "html":
                    preview_format = "html"
                elif format in ["docx", "pptx"]:
                    # Prefer using DocumentGenerator's .preview.html file
                    # (generated directly from research_result, not reverse-engineered from DOCX)
                    preview_html_path = Path(document_path).with_suffix('.preview.html')
                    if preview_html_path.exists():
                        preview_path = str(preview_html_path)
                        preview_format = "html"
                        logger.info(f"Using pre-generated HTML preview: {preview_path}")
                    else:
                        # Fallback: reverse-engineer HTML preview from DOCX
                        html_preview = self._generate_html_from_docx(document_path)
                        if html_preview and os.path.exists(html_preview):
                            preview_path = html_preview
                            preview_format = "html"
                        else:
                            preview_format = "png"  # Fallback to PNG
                else:
                    preview_format = "png"
                
                return PreviewResult(
                    success=True,
                    preview_path=preview_path,
                    preview_format=preview_format,
                    width=width,
                    height=height,
                    file_size=file_size,
                    cached=False
                )
            else:
                return PreviewResult(
                    success=False,
                    error="Preview generation failed",
                    error_code="GENERATION_FAILED"
                )
                
        except (OSError, IOError) as e:
            logger.error(f"Preview generation error: {e}")
            return PreviewResult(
                success=False,
                error=f"File error: {e}",
                error_code="FILE_ERROR"
            )
    
    def generate_html_preview_from_data(
        self,
        research_result: Dict[str, Any],
        output_format: str = "docx"
    ) -> PreviewResult:
        """
        Generate HTML preview directly from research data (without DOCX reverse-engineering).
        
        This is the recommended preview approach - HTML preview is generated directly
        from research_result via ContentOrchestrator, avoiding information loss and
        formatting issues from reverse-engineering HTML from DOCX.
        
        Args:
            research_result: Research result data (includes sections, etc.)
            output_format: Output format (for template selection)
            
        Returns:
            PreviewResult preview result
        """
        try:
            from src.content.content_orchestrator import ContentOrchestrator
            
            # Use ContentOrchestrator to generate HTML
            orchestrator = ContentOrchestrator()
            html_content = orchestrator.transform_to_html(
                research_result=research_result,
                output_format=output_format
            )
            
            # Generate cache key from content (not memory address)
            import json
            try:
                _content_str = json.dumps(research_result, sort_keys=True, ensure_ascii=False, default=str)
            except Exception:
                _content_str = str(len(str(research_result))) + str(hash(str(research_result)) % (2**31))
            cache_key = hashlib.md5(
                _content_str.encode("utf-8"),
                usedforsecurity=False
            ).hexdigest()
            
            # Check cache
            with self._lock:
                if cache_key in self._cache_index:
                    cached_path = self._cache_index[cache_key]
                    if os.path.exists(cached_path):
                        return PreviewResult(
                            success=True,
                            preview_path=cached_path,
                            preview_format="html",
                            cached=True,
                            file_size=os.path.getsize(cached_path)
                        )
            
            # Save HTML file
            task_id = research_result.get("task_id", cache_key)
            html_path = str(self.cache_dir / f"{task_id}_preview.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            file_size = os.path.getsize(html_path)
            
            # Update cache
            with self._lock:
                self._cache_index[cache_key] = html_path
            self._save_cache_index()
            
            logger.info(f"Generated HTML preview from data: {html_path}")
            return PreviewResult(
                success=True,
                preview_path=html_path,
                preview_format="html",
                width=DEFAULT_PREVIEW_WIDTH,
                height=DEFAULT_PREVIEW_HEIGHT,
                file_size=file_size,
                cached=False
            )
            
        except Exception as e:
            logger.error(f"Failed to generate HTML preview from data: {e}")
            return PreviewResult(
                success=False,
                error=f"HTML preview generation failed: {e}",
                error_code="HTML_PREVIEW_FAILED"
            )
    
    def _generate_pdf_preview(
        self,
        document_path: str,
        width: int,
        height: int,
        page_number: int
    ) -> Optional[str]:
        """Generate PDF preview"""
        if self._pdf2image_available:
            try:
                from pdf2image import convert_from_path
                
                images = convert_from_path(
                    document_path,
                    first_page=page_number,
                    last_page=page_number,
                    size=(width, height)
                )
                
                if images:
                    preview_path = str(self.cache_dir / f"{os.path.basename(document_path)}_preview.png")
                    images[0].save(preview_path, 'PNG')
                    return preview_path
                    
            except Exception as e:
                logger.warning(f"pdf2image failed: {e}")
        
        # Fallback: generate placeholder
        return self._generate_placeholder_preview(document_path, width, height)
    
    def _generate_office_preview(
        self,
        document_path: str,
        format: str,
        width: int,
        height: int,
        page_number: int
    ) -> Optional[str]:
        """Generate Office document preview (fallback)"""
        # Simplified: generate placeholder preview
        return self._generate_placeholder_preview(document_path, width, height)
    
    def _generate_html_preview(
        self,
        document_path: str,
        width: int,
        height: int
    ) -> Optional[str]:
        """Generate HTML preview (fallback)"""
        return self._generate_placeholder_preview(document_path, width, height)
    
    def _generate_placeholder_preview(
        self,
        document_path: str,
        width: int,
        height: int
    ) -> Optional[str]:
        """Generate placeholder preview image"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create placeholder image
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Draw border
            draw.rectangle([10, 10, width-10, height-10], outline='gray', width=2)
            
            # Add filename
            filename = os.path.basename(document_path)
            draw.text((20, 20), f"Preview: {filename}", fill='black')
            draw.text((20, 50), f"Size: {width}x{height}", fill='gray')
            
            # Save
            preview_path = str(self.cache_dir / f"{filename}_preview.png")
            img.save(preview_path, 'PNG')
            
            return preview_path
            
        except ImportError:
            # PIL not available, create empty file as marker
            preview_path = str(self.cache_dir / f"{os.path.basename(document_path)}_preview.txt")
            with open(preview_path, 'w', encoding='utf-8') as f:
                f.write(f"Preview placeholder for {document_path}\n")
            return preview_path
    
    def clear_cache(self) -> None:
        """Clear cache"""
        # Delete cache files
        for file in self.cache_dir.iterdir():
            if file.is_file() and file.suffix in ['.png', '.txt']:
                try:
                    file.unlink()
                except OSError:
                    pass
        
        # Clear index
        self._cache_index = {}
        self._save_cache_index()
        
        logger.info("Preview cache cleared")
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics"""
        cached_files = []
        total_size = 0
        
        for file in self.cache_dir.iterdir():
            if file.is_file() and file.suffix in ['.png', '.txt']:
                size = file.stat().st_size
                cached_files.append({
                    "path": str(file),
                    "size": size,
                    "mtime": datetime.fromtimestamp(file.stat().st_mtime).isoformat()
                })
                total_size += size
        
        return {
            "cached_files": cached_files,
            "total_count": len(cached_files),
            "total_size": total_size
        }
    
    def _generate_html_from_docx(self, document_path: str) -> Optional[str]:
        """
        Generate HTML preview from docx document
        
        Args:
            document_path: Word document path
            
        Returns:
            HTML preview file path, None on failure
        """
        try:
            from docx import Document
            
            # Read document
            doc = Document(document_path)
            
            # Generate HTML content
            html_parts = ['<!DOCTYPE html>', '<html>', '<head>', 
                         '<meta charset="utf-8">', 
                         '<title>Preview</title>',
                         '<style>',
                         'body { font-family: "Microsoft YaHei", Arial, sans-serif; margin: 40px; line-height: 1.6; }',
                         'h1 { color: #2C3E50; border-bottom: 2px solid #C9A227; padding-bottom: 10px; }',
                         'h2 { color: #2C3E50; margin-top: 30px; }',
                         'h3 { color: #34495E; }',
                         'p { margin: 10px 0; }',
                         'table { border-collapse: collapse; width: 100%; margin: 20px 0; }',
                         'th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }',
                         'th { background-color: #2C3E50; color: white; }',
                         'img { max-width: 100%; height: auto; }',
                         '</style>',
                         '</head>', '<body>']
            
            # Extract document content
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # Determine heading level by style
                style_name = para.style.name.lower() if para.style else ''
                if 'heading 1' in style_name or 'title' in style_name:
                    html_parts.append(f'<h1>{text}</h1>')
                elif 'heading 2' in style_name:
                    html_parts.append(f'<h2>{text}</h2>')
                elif 'heading 3' in style_name:
                    html_parts.append(f'<h3>{text}</h3>')
                else:
                    html_parts.append(f'<p>{text}</p>')
            
            # Extract tables
            for table in doc.tables:
                html_parts.append('<table>')
                for i, row in enumerate(table.rows):
                    tag = 'th' if i == 0 else 'td'
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            
            html_parts.extend(['</body>', '</html>'])
            
            # Save HTML file
            html_path = str(self.cache_dir / f"{os.path.basename(document_path)}_preview.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_parts))
            
            return html_path
            
        except Exception as e:
            logger.warning(f"Failed to generate HTML preview: {e}")
            return None


# Export
__all__ = ["PreviewGenerator", "PreviewResult"]