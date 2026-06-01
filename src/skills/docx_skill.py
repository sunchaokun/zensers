"""
DocxSkill - Word Document Generation Skill

Generates professional format Word reports, supporting titles, paragraphs, tables, and styles.
"""
from pathlib import Path
from typing import Any, Dict, List, Optional

from src.skills.base import Skill, SkillConfig


class DocxSkill(Skill):
    """
    Word Document Generation Skill

    Supported operations:
    - create: Create new document
    - add_heading: Add heading (levels 1-6)
    - add_paragraph: Add paragraph
    - add_table: Add table
    - build_report: Build complete report in one go
    """

    @property
    def name(self) -> str:
        return "docx_skill"

    @property
    def description(self) -> str:
        return "Word document generation, supports headings/paragraphs/tables, generates professional format reports"

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """
        Execute document operation

        Args:
            action: Operation type
            filepath: Document path
            text/data/title/sections: Content parameters for each operation
        """
        action = kwargs.get("action", "")
        filepath = kwargs.get("filepath", "")

        handlers = {
            "create": self._create,
            "add_heading": self._add_heading,
            "add_paragraph": self._add_paragraph,
            "add_table": self._add_table,
            "build_report": self._build_report,
        }

        handler = handlers.get(action)
        if handler is None:
            return self._failure(f"Unsupported operation: {action}")

        try:
            return await handler(filepath, kwargs)
        except Exception as e:
            return self._failure(str(e), "Document operation failed")

    async def _create(self, filepath: str, kwargs: Dict) -> Dict[str, Any]:
        """Create new document"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set document properties
        title = kwargs.get("title", "")
        author = kwargs.get("author", "Zensers")
        if title:
            doc.core_properties.title = title
        doc.core_properties.author = author

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        doc.save(filepath)
        return self._success({"filepath": filepath}, "Document created successfully")

    async def _add_heading(self, filepath: str, kwargs: Dict) -> Dict[str, Any]:
        """Add heading"""
        from docx import Document

        text = kwargs.get("text", "")
        level = kwargs.get("level", 1)

        doc = Document(filepath)
        doc.add_heading(text, level=level)
        doc.save(filepath)
        return self._success({"filepath": filepath, "heading": text}, "Heading added successfully")

    async def _add_paragraph(self, filepath: str, kwargs: Dict) -> Dict[str, Any]:
        """Add paragraph"""
        from docx import Document

        text = kwargs.get("text", "")
        bold = kwargs.get("bold", False)
        italic = kwargs.get("italic", False)

        doc = Document(filepath)
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        doc.save(filepath)
        return self._success({"filepath": filepath}, "Paragraph added successfully")

    async def _add_table(self, filepath: str, kwargs: Dict) -> Dict[str, Any]:
        """Add table"""
        from docx import Document
        from docx.shared import Pt, RGBColor

        data: List[List[str]] = kwargs.get("data", [])
        if not data:
            return self._failure("Table data cannot be empty")

        doc = Document(filepath)
        rows = len(data)
        cols = max(len(row) for row in data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(data):
            for c_idx, cell_text in enumerate(row):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = str(cell_text)
                # Bold first row
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        doc.save(filepath)
        return self._success(
            {"filepath": filepath, "rows": rows, "cols": cols},
            "Table added successfully"
        )

    async def _build_report(self, filepath: str, kwargs: Dict) -> Dict[str, Any]:
        """Build complete report in one go"""
        title = kwargs.get("title", "Research Report")
        sections = kwargs.get("sections", [])

        # First create document
        create_result = await self._create(filepath, {"title": title})
        if not create_result["success"]:
            return create_result

        # Add cover title
        from docx import Document
        doc = Document(filepath)
        doc.add_heading(title, level=0)
        doc.save(filepath)

        # Add content section by section
        for section in sections:
            heading = section.get("heading", "")
            level = section.get("level", 1)
            content = section.get("content", "")
            table_data = section.get("table")

            if heading:
                await self._add_heading(filepath, {"text": heading, "level": level})

            if content:
                await self._add_paragraph(filepath, {"text": content})

            if table_data:
                await self._add_table(filepath, {"data": table_data})

        return self._success(
            {"filepath": filepath, "sections": len(sections)},
            f"Report built successfully with {len(sections)} sections"
        )
